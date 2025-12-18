"""
High-Performance RAG Pipeline for Technical Documents
Optimized for GPU rental with bge-large-en-v1.5 and re-ranking
Enhanced with non-text content extraction (tables, images, diagrams)
"""

import warnings
# Suppress annoying Pydantic warnings
warnings.filterwarnings("ignore", message=".*validate_default.*")
warnings.filterwarnings("ignore", category=UserWarning, module="pydantic")
# Suppress pypdf warnings for malformed PDFs
warnings.filterwarnings("ignore", message=".*wrong pointing object.*")

import os
import sys
from pathlib import Path

# Add parent directory to Python path so we can import 'backend' module
# This allows the script to be run from any directory
script_dir = Path(__file__).resolve().parent
parent_dir = script_dir.parent
if str(parent_dir) not in sys.path:
    sys.path.insert(0, str(parent_dir))

# Repo-root defaults (makes ingest resilient to being run from backend/ vs repo root)
REPO_ROOT = parent_dir
DEFAULT_DATA_DIR = str(REPO_ROOT / "data")
DEFAULT_STORAGE_DIR = str(REPO_ROOT / "latest_model")

# Set ingestion-safe mode by default (disable metadata updates)
os.environ["DISABLE_METADATA_UPDATE"] = os.environ.get("DISABLE_METADATA_UPDATE", "1")

import logging
import time
import json
import yaml
import re
from typing import List, Dict, Any, Optional, Tuple
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path
import base64
from io import BytesIO

import fitz  # PyMuPDF
import pandas as pd
from PIL import Image
import numpy as np
from tqdm import tqdm

from llama_index.core import SimpleDirectoryReader, VectorStoreIndex, StorageContext, load_index_from_storage, Settings
from llama_index.core.node_parser import SentenceSplitter
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from llama_index.core.schema import NodeWithScore, TextNode, ImageNode, Document
from backend.utils.embedding_utils import build_offline_embedding

# DOCX and Markdown support
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
from llama_index.vector_stores.qdrant import QdrantVectorStore
from llama_index.core.storage.docstore import SimpleDocumentStore
from llama_index.core.storage.index_store import SimpleIndexStore
from sentence_transformers import CrossEncoder
import qdrant_client
import shutil
import tarfile

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Try to import Anthropic (optional dependency)
try:
    from anthropic import Anthropic
    ANTHROPIC_AVAILABLE = True
except ImportError:
    ANTHROPIC_AVAILABLE = False
    logger.warning("Anthropic package not available. Claude rewriting will be disabled.")

# Suppress pypdf warnings for malformed PDFs
logging.getLogger("pypdf").setLevel(logging.ERROR)
logging.getLogger("pypdf._reader").setLevel(logging.ERROR)

# Log DOCX availability after logger is initialized
if not DOCX_AVAILABLE:
    logger.warning("python-docx not available. DOCX support will be disabled.")


class TextPreprocessor:
    """
    Enhanced AI-powered text preprocessor for RAG pipeline.
    Removes boilerplate, normalizes technical content, fixes artifacts, and filters low-quality chunks.
    """
    
    def __init__(self):
        # Compile regex patterns for common boilerplate text
        # Using more specific patterns to avoid catastrophic backtracking
        self.boilerplate_patterns = [
            # Page numbers (various formats) - more specific to avoid hanging
            re.compile(r'^[ \t]*[Pp]age[ \t]+\d+[ \t]*$', re.MULTILINE),
            re.compile(r'^[ \t]*\d{1,4}[ \t]*$', re.MULTILINE),  # Standalone page numbers (limit digits)
            re.compile(r'^[ \t]*-[ \t]*\d+[ \t]*-[ \t]*$', re.MULTILINE),  # "- 5 -"
            
            # Confidential/Proprietary markings
            re.compile(r'\b[Mm]emjet[ \t]+[Cc]onfidential\b', re.IGNORECASE),
            re.compile(r'\b[Cc]onfidential\b', re.IGNORECASE),
            re.compile(r'\b[Pp]roprietary\b', re.IGNORECASE),
            
            # Copyright notices (various formats) - limit length to prevent hanging
            re.compile(r'©[ \t]*\d{4}[^\n]{0,100}$', re.MULTILINE | re.IGNORECASE),
            re.compile(r'Copyright[ \t]+©?[ \t]*\d{4}[^\n]{0,100}$', re.MULTILINE | re.IGNORECASE),
            re.compile(r'All[ \t]+rights[ \t]+reserved\.?', re.IGNORECASE),
            
            # Common header/footer patterns
            re.compile(r'^[ \t]*(Document|Version|Rev|Revision)[ \t]*:[ \t]*[\w\.-]+[ \t]*$', re.MULTILINE | re.IGNORECASE),
            
            # Repeated section titles (if they appear multiple times)
            re.compile(r'^[ \t]*(Table[ \t]+of[ \t]+Contents|Contents|Index)[ \t]*$', re.MULTILINE | re.IGNORECASE),
            
            # Date stamps in headers/footers
            re.compile(r'^[ \t]*\d{1,2}[/-]\d{1,2}[/-]\d{2,4}[ \t]*$', re.MULTILINE),
        ]
        
        # Enhanced header/footer patterns (more comprehensive)
        self.header_footer_patterns = [
            re.compile(r'^[ \t]*(DuraFlex|DuraCore|DuraBolt|anyCUT|EZCut)[ \t]+.*?[ \t]*$', re.MULTILINE | re.IGNORECASE),
            re.compile(r'^[ \t]*[A-Z][a-z]+[ \t]+(Manual|Guide|Databook|Release Notes)[ \t]*$', re.MULTILINE),
            re.compile(r'^[ \t]*V\d+\.\d+[ \t]*$', re.MULTILINE),  # Version numbers
            re.compile(r'^[ \t]*Rev[ \t]*\d+[ \t]*$', re.MULTILINE | re.IGNORECASE),
        ]
        
        # Table of Contents detection patterns
        self.toc_patterns = [
            re.compile(r'^[ \t]*(Table[ \t]+of[ \t]+Contents|Contents|Index|TOC)[ \t]*$', re.MULTILINE | re.IGNORECASE),
            re.compile(r'^\s*\d+\.\d+[ \t]+.*?\s+\d+$', re.MULTILINE),  # "1.2 Section Name    5"
            re.compile(r'^\s*[A-Z][a-z]+[ \t]+\.{3,}[ \t]+\d+$', re.MULTILINE),  # "Section .......... 10"
        ]
        
        # Patterns for structured lines that should be preserved (for smart chunking)
        self.preserve_patterns = [
            re.compile(r'^(Usage|Command|Example|Syntax|Parameters?|Options?|Steps?|Procedure|Note|Warning|Important):\s*', re.MULTILINE | re.IGNORECASE),
            re.compile(r'^\s*\d+[\.\)]\s+', re.MULTILINE),  # Numbered lists: "1. ", "2) "
            re.compile(r'^[-*•]\s+', re.MULTILINE),  # Bullet points
            re.compile(r'^\s*[A-Z][a-z]+:\s*$', re.MULTILINE),  # Section headers ending with colon
        ]
        
        # Common redundant phrases in technical docs
        self.redundant_phrases = [
            (re.compile(r'\bplease[ \t]+note[ \t]+that\b', re.IGNORECASE), ''),
            (re.compile(r'\bit[ \t]+is[ \t]+important[ \t]+to[ \t]+note[ \t]+that\b', re.IGNORECASE), ''),
            (re.compile(r'\bas[ \t]+you[ \t]+can[ \t]+see\b', re.IGNORECASE), ''),
            (re.compile(r'\bas[ \t]+shown[ \t]+above\b', re.IGNORECASE), ''),
            (re.compile(r'\bas[ \t]+shown[ \t]+below\b', re.IGNORECASE), ''),
            (re.compile(r'\bas[ \t]+mentioned[ \t]+previously\b', re.IGNORECASE), ''),
            (re.compile(r'\bas[ \t]+mentioned[ \t]+earlier\b', re.IGNORECASE), ''),
            (re.compile(r'\bfor[ \t]+more[ \t]+information[ \t]+please[ \t]+refer[ \t]+to\b', re.IGNORECASE), 'See'),
            (re.compile(r'\bfor[ \t]+additional[ \t]+details[ \t]+please[ \t]+see\b', re.IGNORECASE), 'See'),
        ]
    
    def is_table_of_contents(self, text: str) -> bool:
        """
        Detect if text is a Table of Contents section.
        Returns True if TOC patterns are found.
        """
        if not text:
            return False
        
        lines = text.split('\n')
        toc_line_count = 0
        
        # Check for TOC header
        for pattern in self.toc_patterns[:1]:  # First pattern is TOC header
            if pattern.search(text):
                toc_line_count += 1
        
        # Check for TOC entry patterns (section numbers with page numbers)
        for line in lines[:20]:  # Check first 20 lines
            for pattern in self.toc_patterns[1:]:
                if pattern.search(line):
                    toc_line_count += 1
                    break
        
        # If we find TOC header + multiple TOC entries, it's likely a TOC
        return toc_line_count >= 3
    
    def remove_table_of_contents(self, text: str) -> str:
        """Remove Table of Contents sections from text."""
        if not self.is_table_of_contents(text):
            return text
        
        lines = text.split('\n')
        cleaned_lines = []
        in_toc = False
        
        for line in lines:
            # Check if line starts TOC
            if any(pattern.search(line) for pattern in self.toc_patterns[:1]):
                in_toc = True
                continue
            
            # Check if line is TOC entry
            if in_toc:
                if any(pattern.search(line) for pattern in self.toc_patterns[1:]):
                    continue
                # Check if we've left TOC (found non-TOC content)
                if line.strip() and not any(pattern.search(line) for pattern in self.toc_patterns):
                    # Look ahead: if next few lines aren't TOC entries, we've left TOC
                    in_toc = False
            
            if not in_toc:
                cleaned_lines.append(line)
        
        return '\n'.join(cleaned_lines)
    
    def remove_headers_footers(self, text: str) -> str:
        """Remove header and footer text that appears on multiple pages."""
        if not text:
            return text
        
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            # Skip header/footer patterns
            is_header_footer = False
            
            # Check enhanced header/footer patterns
            for pattern in self.header_footer_patterns:
                if pattern.search(line):
                    is_header_footer = True
                    break
            
            # Check if line is very short and appears to be header/footer
            if not is_header_footer and len(line.strip()) < 50:
                # Check for common header/footer indicators
                if (line.strip().count(' ') < 5 and 
                    (line.strip().isupper() or 
                     re.match(r'^[A-Z][a-z]+[ \t]+(Manual|Guide|V\d+|Rev)', line.strip()))):
                    is_header_footer = True
            
            if not is_header_footer:
                cleaned_lines.append(line)
        
        return '\n'.join(cleaned_lines)
    
    def is_first_page_without_content(self, text: str, metadata: dict = None) -> bool:
        """
        Detect if this is a first page (cover page) with no meaningful content.
        Checks for title pages, cover pages, etc.
        """
        if not text:
            return True
        
        # Check page number
        page_label = metadata.get('page_label', '') if metadata else ''
        if page_label and page_label not in ['1', 'i', 'I']:
            return False
        
        # Check word count
        words = len(text.split())
        if words < 15:
            return True
        
        # Check for cover page indicators
        cover_indicators = [
            r'^[A-Z][A-Z\s]{10,}$',  # All caps title
            r'^(User|Installation|Service|Operation)[ \t]+(Manual|Guide|Databook)',  # Title format
        ]
        
        lines = text.split('\n')[:10]  # Check first 10 lines
        cover_line_count = 0
        
        for line in lines:
            for pattern in cover_indicators:
                if re.match(pattern, line.strip(), re.IGNORECASE):
                    cover_line_count += 1
        
        # If mostly cover page content and low word count
        return cover_line_count >= 2 and words < 50
    
    def fix_hyphenation(self, text: str) -> str:
        """
        Fix hyphenated words split across lines.
        Example: "print-\nhead" -> "printhead"
        """
        if not text:
            return text
        
        # Pattern: word ending with hyphen, followed by newline, followed by word continuation
        # Match: "word-\nword" -> "wordword"
        text = re.sub(r'([a-zA-Z])-\s*\n\s*([a-zA-Z])', r'\1\2', text)
        
        # Also handle cases with spaces: "word- \n word"
        text = re.sub(r'([a-zA-Z])-\s+\n\s+([a-zA-Z])', r'\1\2', text)
        
        return text
    
    def fix_line_breaks(self, text: str) -> str:
        """
        Fix inappropriate line breaks in the middle of sentences.
        Preserves intentional paragraph breaks (double newlines).
        """
        if not text:
            return text
        
        lines = text.split('\n')
        fixed_lines = []
        i = 0
        
        while i < len(lines):
            line = lines[i].strip()
            
            if not line:
                fixed_lines.append('')
                i += 1
                continue
            
            # Check if line ends with sentence-ending punctuation
            ends_with_punctuation = re.search(r'[.!?]\s*$', line)
            
            # Check if next line starts with capital letter (new sentence)
            next_starts_sentence = False
            if i + 1 < len(lines):
                next_line = lines[i + 1].strip()
                next_starts_sentence = bool(re.match(r'^[A-Z]', next_line))
            
            # If line doesn't end with punctuation and next doesn't start sentence,
            # it's likely a broken line - join them
            if not ends_with_punctuation and not next_starts_sentence and i + 1 < len(lines):
                next_line = lines[i + 1].strip()
                if next_line and not next_line.startswith(('•', '-', '*', '1.', '2.', '3.')):
                    # Join with space
                    line = line + ' ' + next_line
                    i += 1  # Skip next line since we joined it
            
            fixed_lines.append(line)
            i += 1
        
        return '\n'.join(fixed_lines)
    
    def remove_repeated_phrases(self, text: str) -> str:
        """Remove redundant phrases that don't add semantic value."""
        cleaned = text
        
        for pattern, replacement in self.redundant_phrases:
            try:
                cleaned = pattern.sub(replacement, cleaned)
            except Exception as e:
                logger.debug(f"Failed to remove redundant phrase: {e}")
                continue
        
        return cleaned
    
    def normalize_technical_content(self, text: str) -> str:
        """
        Normalize technical instructions and explanations.
        Fixes spacing, punctuation, and formatting issues.
        """
        if not text:
            return text
        
        # Fix spacing around punctuation
        text = re.sub(r'\s+([.,;:!?])', r'\1', text)  # Remove space before punctuation
        text = re.sub(r'([.,;:!?])([^\s])', r'\1 \2', text)  # Add space after punctuation
        
        # Fix spacing around parentheses and brackets
        text = re.sub(r'\(\s+', '(', text)
        text = re.sub(r'\s+\)', ')', text)
        text = re.sub(r'\[\s+', '[', text)
        text = re.sub(r'\s+\]', ']', text)
        
        # Normalize multiple spaces (but preserve intentional spacing in tables/code)
        # Only normalize if not in a table-like structure
        if '|' not in text and '\t' not in text:
            text = re.sub(r' {2,}', ' ', text)
        
        # Fix common technical formatting issues
        text = re.sub(r'(\d+)\s*-\s*(\d+)', r'\1-\2', text)  # Number ranges: "5 - 10" -> "5-10"
        text = re.sub(r'(\w+)\s*/\s*(\w+)', r'\1/\2', text)  # Slashes: "A / B" -> "A/B"
        
        return text
    
    def remove_boilerplate(self, text: str) -> str:
        """Remove common boilerplate text patterns from the input."""
        cleaned = text
        
        # Apply each boilerplate pattern with error handling
        for pattern in self.boilerplate_patterns:
            try:
                cleaned = pattern.sub('', cleaned)
            except Exception as e:
                # If regex fails, skip this pattern and continue
                logger.debug(f"Regex pattern failed, skipping: {e}")
                continue
        
        return cleaned
    
    def normalize_whitespace(self, text: str) -> str:
        """Normalize whitespace: collapse multiple spaces/tabs, normalize newlines."""
        try:
            # Replace tabs with spaces
            text = text.replace('\t', ' ')
            
            # Collapse multiple spaces into single space (limit to prevent hanging)
            text = re.sub(r' {2,}', ' ', text)
            
            # Normalize line breaks: multiple newlines -> double newline (paragraph break)
            text = re.sub(r'\n{3,}', '\n\n', text)
            
            # Remove leading/trailing whitespace from each line
            lines = [line.strip() for line in text.split('\n')]
            text = '\n'.join(lines)
            
            # Remove leading/trailing whitespace from entire text
            text = text.strip()
        except Exception as e:
            logger.debug(f"Whitespace normalization failed: {e}")
            # Return original text if normalization fails
            pass
        
        return text
    
    def should_preserve_line(self, line: str) -> bool:
        """Check if a line matches patterns that should be preserved as-is."""
        for pattern in self.preserve_patterns:
            if pattern.search(line):
                return True
        return False
    
    def clean_text(self, text: str, metadata: dict = None) -> str:
        """
        Apply all enhanced cleaning steps in optimal order.
        Returns cleaned text ready for chunking and embedding.
        """
        if not text:
            return text
        
        cleaned = text
        
        # Step 1: Remove Table of Contents
        cleaned = self.remove_table_of_contents(cleaned)
        
        # Step 2: Remove headers and footers
        cleaned = self.remove_headers_footers(cleaned)
        
        # Step 3: Remove boilerplate (copyright, page numbers, etc.)
        cleaned = self.remove_boilerplate(cleaned)
        
        # Step 4: Fix text artifacts
        cleaned = self.fix_hyphenation(cleaned)
        cleaned = self.fix_line_breaks(cleaned)
        
        # Step 5: Remove redundant phrases
        cleaned = self.remove_repeated_phrases(cleaned)
        
        # Step 6: Normalize technical content (spacing, punctuation)
        cleaned = self.normalize_technical_content(cleaned)
        
        # Step 7: Normalize whitespace (final pass)
        cleaned = self.normalize_whitespace(cleaned)
        
        return cleaned
    
    def is_low_content_page(self, text: str, min_words: int = 15) -> bool:
        """Check if a page has too little content to be useful."""
        words = len(text.split())
        return words < min_words
    
    def should_skip_node(self, text: str, min_chars: int = 30, metadata: dict = None) -> Tuple[bool, str]:
        """
        Check if a node should be skipped (too short or empty).
        Returns (should_skip: bool, reason: str)
        """
        if not text:
            return True, "empty_text"
        
        text_stripped = text.strip()
        if len(text_stripped) < min_chars:
            return True, "too_short"
        
        # Check if it's mostly whitespace or special characters
        alpha_chars = len(re.findall(r'[a-zA-Z]', text))
        if alpha_chars < min_chars // 2:  # At least half should be alphabetic
            return True, "low_alphabetic_content"
        
        # Check if it's a Table of Contents
        if self.is_table_of_contents(text):
            return True, "table_of_contents"
        
        # Check if it's a first page without content
        if self.is_first_page_without_content(text, metadata):
            return True, "first_page_no_content"
        
        return False, ""


class ClaudeSemanticRewriter:
    """
    Uses Claude API to semantically rewrite text chunks for improved clarity
    while preserving technical meaning and structured content.
    """
    
    def __init__(self, api_key: Optional[str] = None, model: str = "claude-3-5-sonnet-20241022", 
                 enabled: bool = False, max_retries: int = 2, timeout: int = 30):
        """
        Initialize Claude semantic rewriter.
        
        Args:
            api_key: Anthropic API key (defaults to ANTHROPIC_API_KEY env var)
            model: Claude model to use (default: claude-3-5-sonnet-20241022)
            enabled: Whether rewriting is enabled (default: False)
            max_retries: Maximum retry attempts per chunk
            timeout: Request timeout in seconds
        """
        self.enabled = enabled and ANTHROPIC_AVAILABLE
        self.model = model
        self.max_retries = max_retries
        self.timeout = timeout
        self.client = None
        
        if self.enabled:
            api_key = api_key or os.getenv('ANTHROPIC_API_KEY')
            if not api_key:
                logger.warning("⚠️ Claude rewriting enabled but ANTHROPIC_API_KEY not found. Disabling rewriting.")
                self.enabled = False
            else:
                try:
                    self.client = Anthropic(api_key=api_key, timeout=timeout)
                    logger.info(f"✅ Claude semantic rewriter initialized (model: {model})")
                except Exception as e:
                    logger.warning(f"⚠️ Failed to initialize Claude client: {e}. Disabling rewriting.")
                    self.enabled = False
    
    def _is_structured_content(self, text: str, metadata: dict) -> bool:
        """
        Check if content is structured (table, code, list) that should be preserved as-is.
        """
        content_type = metadata.get("content_type", "text")
        
        # Don't rewrite tables, images, or captions
        if content_type in ["table", "image", "figure_caption"]:
            return True
        
        # Check for code blocks
        if '```' in text or '`' in text:
            return True
        
        # Check for markdown tables
        if text.count('|') >= 6 and re.search(r'\|[^\|]+\|', text):
            return True
        
        # Check for dense lists (many bullets/numbers)
        lines = text.split('\n')
        list_lines = sum(1 for line in lines if re.match(r'^\s*[-*•\d]', line))
        if list_lines >= len(lines) * 0.5:  # 50% or more are list items
            return True
        
        return False
    
    def _create_rewrite_prompt(self, text: str, metadata: dict) -> str:
        """Create prompt for Claude to rewrite the text."""
        content_type = metadata.get("content_type", "text")
        file_name = metadata.get("file_name", "document")
        
        prompt = f"""Rewrite the following technical documentation text to improve semantic clarity while preserving all technical meaning and accuracy.

**Requirements:**
1. Keep all technical terms, specifications, and measurements exactly as written
2. Remove minor redundancies, filler phrases, and ambiguous wording
3. Improve sentence flow and clarity
4. Preserve all structured content (tables, lists, code blocks) exactly as-is
5. Maintain the same level of technical detail
6. Do not add new information or remove important details
7. Output ONLY the rewritten text, no explanations or markdown formatting

**Source:** {file_name}
**Content Type:** {content_type}

**Original Text:**
{text}

**Rewritten Text:**"""
        
        return prompt
    
    def rewrite_chunk(self, node: TextNode) -> Tuple[TextNode, bool]:
        """
        Rewrite a single chunk using Claude API.
        
        Args:
            node: TextNode to rewrite
            
        Returns:
            Tuple of (rewritten_node, was_rewritten: bool)
        """
        if not self.enabled or not self.client:
            return node, False
        
        # Skip structured content
        if self._is_structured_content(node.text, node.metadata):
            logger.debug(f"Skipping structured content rewrite: {node.metadata.get('file_name', 'unknown')}")
            return node, False
        
        # Skip very short chunks (not worth API call)
        if len(node.text.strip()) < 100:
            return node, False
        
        # Create prompt
        prompt = self._create_rewrite_prompt(node.text, node.metadata)
        
        # Call Claude API with retries
        for attempt in range(self.max_retries):
            try:
                response = self.client.messages.create(
                    model=self.model,
                    max_tokens=4096,
                    temperature=0.1,  # Low temperature for consistency
                    messages=[{
                        "role": "user",
                        "content": prompt
                    }]
                )
                
                # Extract rewritten text
                rewritten_text = response.content[0].text.strip()
                
                # Validate: rewritten text should not be empty and should be reasonable length
                if not rewritten_text or len(rewritten_text) < len(node.text) * 0.5:
                    logger.warning(f"Claude rewrite too short or empty, using original. Chunk: {node.metadata.get('file_name', 'unknown')}")
                    return node, False
                
                # Create new node with rewritten text but same metadata
                rewritten_node = TextNode(
                    text=rewritten_text,
                    metadata=node.metadata.copy()  # Preserve all metadata
                )
                
                logger.debug(f"Successfully rewritten chunk from {node.metadata.get('file_name', 'unknown')}")
                return rewritten_node, True
                
            except Exception as e:
                if attempt < self.max_retries - 1:
                    logger.warning(f"Claude rewrite attempt {attempt + 1} failed, retrying: {e}")
                    time.sleep(1)  # Brief delay before retry
                else:
                    logger.warning(f"Claude rewrite failed after {self.max_retries} attempts, using original: {e}")
                    return node, False
        
        return node, False
    
    def rewrite_nodes(self, nodes: List[TextNode], show_progress: bool = True) -> Tuple[List[TextNode], Dict[str, int]]:
        """
        Rewrite a list of TextNodes using Claude API.
        
        Args:
            nodes: List of TextNode objects to rewrite
            show_progress: Whether to show progress bar
            
        Returns:
            Tuple of (rewritten_nodes, stats_dict)
        """
        if not self.enabled:
            logger.info("Claude rewriting is disabled, skipping rewrite step")
            return nodes, {"rewritten": 0, "skipped": len(nodes), "failed": 0, "structured": 0}
        
        logger.info(f"🔄 Starting Claude semantic rewriting for {len(nodes)} chunks...")
        
        rewritten_nodes = []
        stats = {
            "rewritten": 0,
            "skipped": 0,
            "failed": 0,
            "structured": 0
        }
        
        # Process nodes with progress bar
        iterator = tqdm(nodes, desc="   Rewriting chunks", disable=not show_progress) if show_progress else nodes
        
        for node in iterator:
            # Check if structured content
            if self._is_structured_content(node.text, node.metadata):
                rewritten_nodes.append(node)
                stats["structured"] += 1
                continue
            
            # Rewrite the chunk
            rewritten_node, was_rewritten = self.rewrite_chunk(node)
            rewritten_nodes.append(rewritten_node)
            
            if was_rewritten:
                stats["rewritten"] += 1
            else:
                # Determine why it wasn't rewritten
                if not self.enabled:
                    stats["skipped"] += 1
                elif len(node.text.strip()) < 100:
                    stats["skipped"] += 1
                else:
                    stats["failed"] += 1
        
        logger.info(f"✅ Claude rewriting complete: {stats['rewritten']} rewritten, {stats['structured']} structured (preserved), "
                   f"{stats['skipped']} skipped, {stats['failed']} failed")
        
        return rewritten_nodes, stats


class SmartChunkSplitter:
    """
    Wrapper around SentenceSplitter that preserves structured content like tables,
    code blocks, numbered steps, and command syntax.
    """
    
    def __init__(self, chunk_size: int = 350, chunk_overlap: int = 88, preprocessor: TextPreprocessor = None):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.base_splitter = SentenceSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            include_metadata=True
        )
        self.preprocessor = preprocessor or TextPreprocessor()
    
    def _is_table_content(self, text: str) -> bool:
        """Detect if text contains table-like content (markdown table or pipe-separated)."""
        # Check for markdown table pattern (| col1 | col2 |)
        if re.search(r'\|[^\|]+\|', text) and text.count('|') >= 6:
            return True
        # Check for tab-separated or multiple spaces (likely table)
        lines = text.split('\n')
        if len(lines) >= 3:
            tabs_or_spaces = sum(1 for line in lines if '\t' in line or re.search(r' {3,}', line))
            if tabs_or_spaces >= len(lines) * 0.6:  # 60% of lines have table-like spacing
                return True
        return False
    
    def _is_code_block(self, text: str) -> bool:
        """Detect if text contains code blocks."""
        # Check for code block markers
        if '```' in text or '`' in text:
            return True
        # Check for common code patterns
        code_keywords = ['def ', 'class ', 'import ', 'function', 'return', 'const ', 'var ', 'let ']
        if any(keyword in text for keyword in code_keywords):
            return True
        return False
    
    def _preserve_structured_chunks(self, text: str) -> List[str]:
        """
        Split text while preserving structured content as single units.
        Returns list of text chunks.
        """
        chunks = []
        lines = text.split('\n')
        current_chunk = []
        current_chunk_size = 0
        
        i = 0
        while i < len(lines):
            line = lines[i]
            line_stripped = line.strip()
            
            # Check if this line starts a structured block
            if self.preprocessor.should_preserve_line(line_stripped):
                # Try to collect the entire structured block
                structured_block = [line]
                block_size = len(line)
                j = i + 1
                
                # Collect lines until we hit a non-structured line or exceed chunk size
                while j < len(lines) and block_size + len(lines[j]) < self.chunk_size:
                    next_line = lines[j].strip()
                    # Continue if it's part of the structure (numbered, bullet, or continuation)
                    if (next_line.startswith((' ', '\t')) or  # Indented continuation
                        re.match(r'^\d+[\.\)]', next_line) or  # Next numbered item
                        re.match(r'^[-*•]', next_line) or  # Next bullet
                        re.match(r'^[A-Z][a-z]+:\s*$', next_line)):  # Next section header
                        structured_block.append(lines[j])
                        block_size += len(lines[j])
                        j += 1
                    else:
                        break
                
                # If we have accumulated text, save it first
                if current_chunk:
                    chunks.append('\n'.join(current_chunk))
                    current_chunk = []
                    current_chunk_size = 0
                
                # Add the structured block as a single chunk
                structured_text = '\n'.join(structured_block)
                # Check if it exceeds chunk size (rare, but handle it)
                if len(structured_text) > self.chunk_size:
                    # Split the structured block using base splitter, but preserve internal structure
                    sub_chunks = self.base_splitter.split_text(structured_text)
                    chunks.extend(sub_chunks)
                else:
                    chunks.append(structured_text)
                
                i = j
                continue
            
            # Regular line - add to current chunk
            line_size = len(line)
            if current_chunk_size + line_size <= self.chunk_size:
                current_chunk.append(line)
                current_chunk_size += line_size
                i += 1
            else:
                # Current chunk is full, save it
                if current_chunk:
                    chunks.append('\n'.join(current_chunk))
                    # Start new chunk with overlap
                    overlap_lines = []
                    overlap_size = 0
                    # Get last few lines for overlap (up to chunk_overlap chars)
                    for line_idx in range(len(current_chunk) - 1, -1, -1):
                        if overlap_size + len(current_chunk[line_idx]) <= self.chunk_overlap:
                            overlap_lines.insert(0, current_chunk[line_idx])
                            overlap_size += len(current_chunk[line_idx])
                        else:
                            break
                    current_chunk = overlap_lines + [line]
                    current_chunk_size = overlap_size + line_size
                    i += 1
        
        # Add remaining chunk
        if current_chunk:
            chunks.append('\n'.join(current_chunk))
        
        return chunks
    
    def split_text(self, text: str) -> List[str]:
        """Split text with smart chunking that preserves structured content."""
        # Check if entire text is a table or code block
        if self._is_table_content(text) or self._is_code_block(text):
            # Preserve as single chunk (may exceed chunk_size, but that's okay for structured content)
            return [text]
        
        # Use smart chunking
        return self._preserve_structured_chunks(text)
    
    def get_nodes_from_documents(self, documents: List[Document], show_progress: bool = False) -> List[TextNode]:
        """
        Split documents into nodes with smart chunking.
        This is a simplified version that processes documents one by one.
        """
        all_nodes = []
        
        for doc in tqdm(documents, desc="Splitting documents", disable=not show_progress):
            try:
                text = doc.text or ""
                doc_name = doc.metadata.get('file_name', 'unknown')
                
                # Clean the text first (with error handling)
                try:
                    text = self.preprocessor.clean_text(text, metadata=doc.metadata)
                except Exception as e:
                    logger.warning(f"Error cleaning text for {doc_name}: {e}")
                    # Use original text if cleaning fails
                    text = doc.text or ""
                
                # Check if page/document should be skipped (low content)
                try:
                    if self.preprocessor.is_low_content_page(text):
                        logger.debug(f"Skipping low-content page: {doc_name}")
                        continue
                except Exception as e:
                    logger.debug(f"Error checking low-content for {doc_name}: {e}")
                    # Continue processing if check fails
                
                # Split into chunks (with error handling)
                try:
                    chunks = self.split_text(text)
                except Exception as e:
                    logger.error(f"Error splitting text for {doc_name}: {e}")
                    # Fallback to simple split if smart chunking fails
                    if text:
                        chunks = [text]
                    else:
                        chunks = []
                
                # Create nodes from chunks
                for chunk_idx, chunk_text in enumerate(chunks):
                    # Enhanced skip check with reason tracking
                    should_skip, skip_reason = self.preprocessor.should_skip_node(
                        chunk_text, 
                        metadata={**doc.metadata, "chunk_index": chunk_idx, "total_chunks": len(chunks)}
                    )
                    
                    if should_skip:
                        logger.debug(f"Skipping chunk {chunk_idx} from {doc_name}: {skip_reason}")
                        continue
                    
                    # Create node with metadata (including skip reason if applicable)
                    try:
                        node = TextNode(
                            text=chunk_text,
                            metadata={
                                **doc.metadata,
                                "chunk_index": chunk_idx,
                                "total_chunks": len(chunks),
                                "content_type": "text"
                            }
                        )
                        all_nodes.append(node)
                    except Exception as e:
                        logger.error(f"Error creating node for {doc_name}, chunk {chunk_idx}: {e}")
                        continue
                        
            except Exception as e:
                logger.error(f"Error processing document {doc.metadata.get('file_name', 'unknown')}: {e}")
                continue
        
        return all_nodes


class NonTextExtractor:
    """Extract and process non-text content from documents."""
    
    def __init__(self, output_dir: Optional[str] = None):
        override_dir = os.getenv("EXTRACTED_CONTENT_DIR")
        if override_dir:
            target_dir = Path(override_dir)
        elif output_dir:
            target_dir = Path(output_dir)
        else:
            target_dir = Path(__file__).resolve().parent.parent / "extracted_content"
        self.output_dir = target_dir
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
    def extract_tables_from_pdf(self, pdf_path: str) -> List[Dict[str, Any]]:
        """Extract tables from PDF using PyMuPDF."""
        tables = []
        doc = fitz.open(pdf_path)
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            # Extract tables using PyMuPDF's table detection
            page_tables = page.find_tables()
            
            for table_idx, table in enumerate(page_tables):
                try:
                    # Extract table data
                    table_data = table.extract()
                    if table_data and len(table_data) > 1:  # Ensure we have headers and data
                        # Convert to pandas DataFrame for better structure
                        headers = table_data[0]
                        # Handle duplicate column names
                        unique_headers = []
                        for i, header in enumerate(headers):
                            if header in unique_headers:
                                unique_headers.append(f"{header}_{i}")
                            else:
                                unique_headers.append(header)
                        
                        df = pd.DataFrame(table_data[1:], columns=unique_headers)
                        
                        # Create table metadata
                        table_info = {
                            "source_path": pdf_path,
                            "page_number": page_num + 1,
                            "table_index": table_idx,
                            "table_data": df.to_dict('records'),
                            "table_markdown": df.to_markdown(index=False),
                            "table_json": df.to_json(orient='records'),
                            "row_count": len(df),
                            "column_count": len(df.columns),
                            "content_type": "table"
                        }
                        tables.append(table_info)
                        
                        # Save table as separate file
                        table_filename = f"{Path(pdf_path).stem}_page{page_num+1}_table{table_idx}.json"
                        table_path = self.output_dir / table_filename
                        with open(table_path, 'w', encoding='utf-8') as f:
                            json.dump(table_info, f, indent=2, ensure_ascii=False)
                            
                except Exception as e:
                    logger.warning(f"Failed to extract table {table_idx} from page {page_num + 1}: {e}")
                    continue
                    
        doc.close()
        return tables
    
    def extract_images_from_pdf(self, pdf_path: str) -> List[Dict[str, Any]]:
        """Extract images and diagrams from PDF."""
        images = []
        doc = fitz.open(pdf_path)
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            image_list = page.get_images()
            
            for img_idx, img in enumerate(image_list):
                try:
                    # Get image data
                    xref = img[0]
                    pix = fitz.Pixmap(doc, xref)
                    
                    if pix.n - pix.alpha < 4:  # GRAY or RGB
                        # Convert to PIL Image
                        img_data = pix.tobytes("png")
                        pil_image = Image.open(BytesIO(img_data))
                        
                        # Get image metadata
                        img_rects = page.get_image_rects(xref)
                        img_rect = img_rects[0] if img_rects else None
                        
                        # Create image info (NO base64 embedding - only metadata)
                        image_info = {
                            "source_path": pdf_path,
                            "page_number": page_num + 1,
                            "image_index": img_idx,
                            # Removed: image_data base64 encoding (not needed for embeddings)
                            "width": pil_image.width,
                            "height": pil_image.height,
                            "format": "PNG",
                            "content_type": "image",
                            "caption": f"Image from {Path(pdf_path).stem}, page {page_num + 1}",
                            "bbox": str(img_rect) if img_rect else None
                        }
                        images.append(image_info)
                        
                        # Save image
                        img_filename = f"{Path(pdf_path).stem}_page{page_num+1}_img{img_idx}.png"
                        img_path = self.output_dir / img_filename
                        pil_image.save(img_path)
                        image_info["saved_path"] = str(img_path)
                        
                except Exception as e:
                    logger.warning(f"Failed to extract image {img_idx} from page {page_num + 1}: {e}")
                    continue
                    
        doc.close()
        return images
    
    def extract_figure_captions(self, pdf_path: str) -> List[Dict[str, Any]]:
        """Extract figure captions and references."""
        captions = []
        doc = fitz.open(pdf_path)
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text()
            
            # Look for figure captions (simple pattern matching)
            lines = text.split('\n')
            for i, line in enumerate(lines):
                line_lower = line.lower().strip()
                if any(keyword in line_lower for keyword in ['figure', 'fig.', 'diagram', 'chart', 'graph']):
                    # Extract caption text
                    caption_text = line.strip()
                    if len(caption_text) > 10:  # Filter out very short matches
                        caption_info = {
                            "source_path": pdf_path,
                            "page_number": page_num + 1,
                            "caption_text": caption_text,
                            "content_type": "figure_caption",
                            "line_number": i + 1
                        }
                        captions.append(caption_info)
        
        doc.close()
        return captions


class DocumentLoader:
    """
    Custom document loader that supports PDF, DOCX, and Markdown files.
    Preserves document provenance with file_name and page_label/section metadata.
    Can load from local directory or GCS bucket.
    """
    
    def __init__(
        self,
        data_dir: str = None,
        gcs_bucket: str = None,
        gcs_prefix: str = None,
        manifest_path: str | None = None,
    ):
        self.data_dir = Path(data_dir) if data_dir else None
        self.gcs_bucket = gcs_bucket
        self.gcs_prefix = gcs_prefix or ""
        self.manifest_path = manifest_path
        self.supported_extensions = {'.pdf', '.docx', '.md', '.markdown'}
        self.temp_files = []  # Track temp files for cleanup
    
    def load_documents(self, use_database: bool = True) -> List[Document]:
        """
        Load all supported documents from database (preferred), GCS bucket, or local directory.
        Returns list of Document objects with proper metadata.
        
        Args:
            use_database: If True, load documents from database (DocumentIngestionMetadata with gcs_path).
                         Only processes documents that exist in the database.
        """
        documents = []

        # Highest priority: explicit manifest (deterministic ingestion inputs)
        if self.manifest_path:
            documents = self._load_from_manifest(self.manifest_path)
            logger.info(f"Loaded {len(documents)} document sections from manifest: {self.manifest_path}")
            return documents
        
        # Priority 1: Load from database (if enabled and GCS is configured)
        if use_database:
            try:
                documents = self._load_from_database()
                if documents:
                    logger.info(f"Loaded {len(documents)} documents from database")
                    return documents
                else:
                    logger.warning("No documents found in database, falling back to GCS/local")
            except Exception as e:
                logger.warning(f"Failed to load from database: {e}, falling back to GCS/local", exc_info=True)
        
        # Priority 2: Load from GCS bucket (if configured)
        if self.gcs_bucket:
            documents = self._load_from_gcs()
        # Priority 3: Load from local directory
        elif self.data_dir:
            documents = self._load_from_local()
        else:
            raise ValueError("Either data_dir or gcs_bucket must be provided, or use_database=True with database records")
        
        return documents

    def _load_from_manifest(self, manifest_path: str) -> List[Document]:
        """
        Load documents from a staging manifest written by the production ingestion flow.

        Manifest format:
            { "documents": [ { document_id, gcs_object_name, local_path, machine_models, ... } ] }

        This loader is responsible for injecting REQUIRED per-chunk metadata via doc.metadata, so that
        the SmartChunkSplitter will propagate them into every node:
          - document_id (MUST be a stable, non-empty string; DB-native id preferred; UUID5 fallback if missing)
          - machine_models (list[str])
          - source_gcs (gs://...)
        """
        with open(manifest_path, "r", encoding="utf-8") as f:
            manifest = json.load(f) or {}
        entries = manifest.get("documents", [])
        if not isinstance(entries, list):
            raise ValueError(f"Invalid manifest format (documents must be a list): {manifest_path}")

        documents: list[Document] = []
        for entry in tqdm(entries, desc="Loading documents from manifest"):
            try:
                local_path = entry.get("local_path")
                if not local_path:
                    logger.warning("Skipping manifest entry with missing local_path", extra={"entry": entry})
                    continue
                file_path = Path(local_path)
                if not file_path.exists():
                    logger.warning("Skipping manifest entry (local file missing)", extra={"local_path": local_path})
                    continue

                filename = entry.get("filename") or file_path.name
                file_ext = file_path.suffix.lower()
                if file_ext not in self.supported_extensions:
                    continue

                # REQUIRED metadata for every chunk/node
                document_id = entry.get("document_id")
                machine_models = entry.get("machine_models") or entry.get("machine_model_names") or []
                machine_model_names = entry.get("machine_model_names") or machine_models or []
                machine_model_ids = entry.get("machine_model_ids") or []
                source_gcs = entry.get("source_gcs") or entry.get("gcs_uri") or entry.get("gcs_path")
                ingestion_metadata_id = entry.get("ingestion_metadata_id") or entry.get("metadata_id")

                # Normalize machine_models
                if isinstance(machine_models, str):
                    try:
                        machine_models = json.loads(machine_models)
                    except Exception:
                        machine_models = [m.strip() for m in machine_models.split(",") if m.strip()]
                if not isinstance(machine_models, list):
                    machine_models = []
                machine_models = [m for m in machine_models if isinstance(m, str) and m.strip()]

                if isinstance(machine_model_names, str):
                    try:
                        machine_model_names = json.loads(machine_model_names)
                    except Exception:
                        machine_model_names = [m.strip() for m in machine_model_names.split(",") if m.strip()]
                if not isinstance(machine_model_names, list):
                    machine_model_names = []
                machine_model_names = [m for m in machine_model_names if isinstance(m, str) and m.strip()]

                if isinstance(machine_model_ids, str):
                    try:
                        machine_model_ids = json.loads(machine_model_ids)
                    except Exception:
                        machine_model_ids = [m.strip() for m in machine_model_ids.split(",") if m.strip()]
                if not isinstance(machine_model_ids, list):
                    machine_model_ids = []
                machine_model_ids = [m for m in machine_model_ids if isinstance(m, str) and m.strip()]

                if document_id is None:
                    logger.warning(
                        "Manifest entry missing document_id; setting to 0 (will break document_id-based deletion)",
                        extra={"filename": filename, "source_gcs": source_gcs},
                    )
                    document_id = "0"
                document_id = str(document_id)

                base_meta = {
                    "file_name": filename,
                    "file_type": file_ext.lstrip(".") if file_ext else "unknown",
                    "gcs_path": source_gcs,  # historical key used elsewhere
                    "source_gcs": source_gcs,
                    "local_path": str(file_path.resolve()),
                    "document_id": document_id,
                    # Best-practice: store both ids + names, plus backwards-compat aliases
                    "machine_model_ids": machine_model_ids,
                    "machine_model_names": machine_model_names or machine_models,
                    "machine_models": machine_model_names or machine_models,
                    # Backwards compatibility: orchestrator uses machine_model (string|list). Use list[str].
                    "machine_model": machine_model_names or machine_models,
                }
                if ingestion_metadata_id:
                    base_meta["ingestion_metadata_id"] = ingestion_metadata_id
                    base_meta["metadata_id"] = ingestion_metadata_id  # legacy alias

                # Load document based on type, and inject base metadata onto every produced section
                if file_ext == ".pdf":
                    pdf_docs = SimpleDirectoryReader(input_files=[str(file_path)]).load_data()
                    for doc in pdf_docs:
                        doc.metadata = {**base_meta, **(doc.metadata or {})}
                    documents.extend(pdf_docs)
                elif file_ext == ".docx" and DOCX_AVAILABLE:
                    docx_docs = self._load_docx(file_path)
                    for doc in docx_docs:
                        doc.metadata = {**base_meta, **(doc.metadata or {})}
                    documents.extend(docx_docs)
                elif file_ext in {".md", ".markdown"}:
                    md_docs = self._load_markdown(file_path)
                    for doc in md_docs:
                        doc.metadata = {**base_meta, **(doc.metadata or {})}
                    documents.extend(md_docs)
            except Exception as e:
                logger.error(f"Error loading manifest entry: {e}", exc_info=True)
                continue

        return documents
    
    def _load_from_database(self) -> List[Document]:
        """
        Load documents from database (DocumentIngestionMetadata table).
        Only processes documents that have gcs_path set in the database.
        """
        from backend.utils.db import SessionLocal, DocumentIngestionMetadata, Document as DBDocument
        from backend.utils.gcs_client import download_to_file
        import tempfile
        
        # Validate database connection matches expected configuration
        from backend.config.env import settings
        db_url = settings.DATABASE_URL if hasattr(settings, 'DATABASE_URL') else os.getenv('DATABASE_URL', 'NOT_SET')
        
        # Fail fast if SQLite detected (should never happen in production)
        if db_url.startswith('sqlite'):
            raise RuntimeError(
                f"❌ CRITICAL: ingest.py detected SQLite database ({db_url[:50]}...). "
                "This should NEVER happen in production. "
                "Ensure DATABASE_URL points to PostgreSQL. "
                "If ENV=prod but using SQLite, this is a configuration error."
            )
        
        session = SessionLocal()
        documents = []
        temp_dir = tempfile.mkdtemp(prefix="ingest_db_")
        self.temp_files.append(temp_dir)  # Track for cleanup
        
        try:
            # Log database connection info (without secrets)
            from backend.config.env import settings
            db_url = settings.DATABASE_URL if hasattr(settings, 'DATABASE_URL') else os.getenv('DATABASE_URL', 'NOT_SET')
            # Mask password in connection string for logging
            import re
            if db_url and db_url != 'NOT_SET':
                db_url_safe = re.sub(r':([^:@]+)@', r':***@', db_url)
                db_host = re.search(r'@([^:/]+)', db_url)
                db_name = re.search(r'/([^?]+)', db_url)
                logger.info(f"🔍 Ingest.py using DATABASE_URL: {db_url_safe}")
                logger.info(f"   Database host: {db_host.group(1) if db_host else 'unknown'}")
                logger.info(f"   Database name: {db_name.group(1) if db_name else 'unknown'}")
            else:
                logger.warning("⚠️ DATABASE_URL not set in ingest.py!")
            
            # Log GCS configuration
            logger.info(f"🔍 GCS Bucket: {self.gcs_bucket}, Prefix: {self.gcs_prefix}")
            
            # Query all DocumentIngestionMetadata records that have gcs_path
            # Join with Document table to get gcs_path
            # Only process documents that are active and have GCS paths
            from sqlalchemy import or_, func
            
            # First, get total counts for validation
            total_metadata_count = session.query(func.count(DocumentIngestionMetadata.id)).scalar() or 0
            total_document_count = session.query(func.count(DBDocument.id)).scalar() or 0
            documents_with_gcs = session.query(func.count(DBDocument.id)).filter(
                DBDocument.gcs_path.isnot(None)
            ).scalar() or 0
            metadata_with_gcs_path = session.query(func.count(DocumentIngestionMetadata.id)).filter(
                DocumentIngestionMetadata.file_path.like('gs://%')
            ).scalar() or 0
            
            logger.info(f"📊 Database counts:")
            logger.info(f"   Total DocumentIngestionMetadata records: {total_metadata_count}")
            logger.info(f"   Total Document records: {total_document_count}")
            logger.info(f"   Document records with gcs_path: {documents_with_gcs}")
            logger.info(f"   Metadata records with gs:// file_path: {metadata_with_gcs_path}")
            
            # Query all DocumentIngestionMetadata records that have gcs_path
            # Join with Document table to get gcs_path
            # Only process documents that are active and have GCS paths
            metadata_records = (
                session.query(DocumentIngestionMetadata, DBDocument)
                .outerjoin(DBDocument, DocumentIngestionMetadata.filename == DBDocument.file_name)
                .filter(
                    # Must have gcs_path (either from Document or from metadata file_path)
                    or_(
                        DBDocument.gcs_path.isnot(None),
                        (DocumentIngestionMetadata.file_path.isnot(None) & 
                         DocumentIngestionMetadata.file_path.like('gs://%'))
                    )
                )
                .filter(
                    # Only process active documents (if Document record exists)
                    or_(
                        DBDocument.is_active.is_(True),
                        DBDocument.id.is_(None)  # No Document record yet, process anyway
                    )
                )
                .all()
            )
            
            logger.info(f"✅ Found {len(metadata_records)} documents in database with GCS paths (matching ingest query)")
            
            # Validation: Check for orphaned records
            # Documents in metadata but not matching query
            all_metadata = session.query(DocumentIngestionMetadata).all()
            matched_filenames = {meta.filename for meta, _ in metadata_records}
            orphaned_metadata = [meta for meta in all_metadata if meta.filename not in matched_filenames]
            
            if orphaned_metadata:
                logger.warning(f"⚠️ Found {len(orphaned_metadata)} DocumentIngestionMetadata records without GCS paths:")
                for meta in orphaned_metadata[:10]:  # Log first 10
                    has_doc = session.query(DBDocument).filter(DBDocument.file_name == meta.filename).first()
                    doc_gcs = has_doc.gcs_path if has_doc else None
                    logger.warning(f"   - {meta.filename}: status={meta.status}, file_path={meta.file_path}, doc_gcs_path={doc_gcs}")
                if len(orphaned_metadata) > 10:
                    logger.warning(f"   ... and {len(orphaned_metadata) - 10} more")
            
            # Validation: Compare with GCS storage (if available)
            if self.gcs_bucket:
                try:
                    from backend.utils.gcs_client import list_object_names
                    gcs_objects = list_object_names(self.gcs_bucket, self.gcs_prefix)
                    gcs_filenames = set()
                    for obj_name in gcs_objects:
                        # Extract filename from GCS path (format: prefix/metadata_id/filename)
                        # Remove prefix to get relative path
                        rel_path = obj_name.replace(self.gcs_prefix, '').lstrip('/')
                        parts = rel_path.split('/')
                        if len(parts) >= 2:
                            # Format: metadata_id/filename - last part is filename
                            gcs_filenames.add(parts[-1])
                    
                    db_filenames = {meta.filename for meta, _ in metadata_records}
                    orphaned_storage = gcs_filenames - db_filenames
                    missing_storage = db_filenames - gcs_filenames
                    
                    logger.info(f"📊 Storage validation:")
                    logger.info(f"   GCS objects found: {len(gcs_objects)}")
                    logger.info(f"   Unique filenames in GCS: {len(gcs_filenames)}")
                    logger.info(f"   Documents in DB with GCS paths: {len(db_filenames)}")
                    
                    if orphaned_storage:
                        logger.warning(f"⚠️ Found {len(orphaned_storage)} ORPHANED STORAGE OBJECTS (in GCS but not in DB):")
                        for filename in list(orphaned_storage)[:5]:
                            logger.warning(f"   - {filename}")
                        if len(orphaned_storage) > 5:
                            logger.warning(f"   ... and {len(orphaned_storage) - 5} more")
                        logger.warning("   These objects exist in GCS but have no matching database record.")
                        logger.warning("   Run with --repair-orphans flag to create DB records (if safe).")
                    
                    if missing_storage:
                        logger.warning(f"⚠️ Found {len(missing_storage)} documents in DB but missing from GCS:")
                        for filename in list(missing_storage)[:5]:
                            logger.warning(f"   - {filename}")
                        if len(missing_storage) > 5:
                            logger.warning(f"   ... and {len(missing_storage) - 5} more")
                        logger.warning("   These database records reference GCS paths that don't exist.")
                            
                except Exception as e:
                    logger.warning(f"⚠️ Could not validate against GCS storage: {e}")
            
            for metadata, db_doc in tqdm(metadata_records, desc="Loading documents from database"):
                try:
                    # Prefer gcs_path from Document table, fallback to file_path from metadata
                    gcs_path = None
                    if db_doc and db_doc.gcs_path:
                        gcs_path = db_doc.gcs_path
                    elif metadata.file_path and metadata.file_path.startswith('gs://'):
                        gcs_path = metadata.file_path
                    
                    if not gcs_path:
                        logger.warning(f"Skipping {metadata.filename}: no GCS path found")
                        continue
                    
                    filename = metadata.filename
                    file_ext = os.path.splitext(filename)[1].lower()
                    
                    # Only process supported file types
                    if file_ext not in self.supported_extensions:
                        logger.debug(f"Skipping {filename}: unsupported file type {file_ext}")
                        continue
                    
                    # Download from GCS to temporary file
                    temp_file_path = os.path.join(temp_dir, filename)
                    
                    logger.debug(f"Downloading {gcs_path} to {temp_file_path}")
                    if not download_to_file(gcs_path, temp_file_path):
                        logger.error(f"Failed to download {gcs_path}")
                        continue
                    
                    self.temp_files.append(temp_file_path)  # Track for cleanup
                    
                    # Load document based on file type
                    file_path = Path(temp_file_path)
                    
                    # Determine document_id and machine models (names + IDs)
                    document_id = db_doc.id if db_doc else None

                    machine_model_names: list[str] = []
                    machine_model_ids: list[int] = []

                    try:
                        if db_doc and hasattr(db_doc, "machine_models") and db_doc.machine_models:
                            machine_model_names = [m.name for m in db_doc.machine_models if getattr(m, "name", None)]
                            machine_model_ids = [int(m.id) for m in db_doc.machine_models if getattr(m, "id", None) is not None]
                    except Exception:
                        machine_model_names = []
                        machine_model_ids = []

                    # Fallback to legacy string fields
                    if not machine_model_names:
                        raw = None
                        if db_doc and db_doc.machine_model:
                            raw = db_doc.machine_model
                        elif metadata.machine_model:
                            raw = metadata.machine_model
                        # Use the shared helper below in this file
                        try:
                            machine_model_names = _parse_machine_models(raw)
                        except Exception:
                            machine_model_names = []

                    # Resolve IDs from names (best-effort)
                    if machine_model_names and not machine_model_ids:
                        try:
                            from backend.utils.db import MachineModel as DBMachineModel
                            rows = session.query(DBMachineModel).filter(DBMachineModel.name.in_(machine_model_names)).all()
                            machine_model_ids = [int(r.id) for r in rows if getattr(r, "id", None) is not None]
                        except Exception:
                            machine_model_ids = []
                    
                    if file_ext == '.pdf':
                        pdf_docs = SimpleDirectoryReader(input_files=[str(file_path)]).load_data()
                        for doc in pdf_docs:
                            doc.metadata['file_name'] = filename
                            doc.metadata['file_type'] = 'pdf'
                            doc.metadata['gcs_path'] = gcs_path
                            doc.metadata['ingestion_metadata_id'] = metadata.id
                            doc.metadata['metadata_id'] = metadata.id  # Keep for backward compatibility
                            doc.metadata['document_id'] = document_id
                            doc.metadata['machine_model_ids'] = machine_model_ids
                            doc.metadata['machine_model_names'] = machine_model_names
                            if machine_model_names:
                                doc.metadata['machine_model'] = machine_model_names
                        documents.extend(pdf_docs)
                    elif file_ext == '.docx' and DOCX_AVAILABLE:
                        docx_docs = self._load_docx(file_path)
                        for doc in docx_docs:
                            doc.metadata['gcs_path'] = gcs_path
                            doc.metadata['ingestion_metadata_id'] = metadata.id
                            doc.metadata['metadata_id'] = metadata.id  # Keep for backward compatibility
                            doc.metadata['document_id'] = document_id
                            doc.metadata['machine_model_ids'] = machine_model_ids
                            doc.metadata['machine_model_names'] = machine_model_names
                            if machine_model_names:
                                doc.metadata['machine_model'] = machine_model_names
                        documents.extend(docx_docs)
                    elif file_ext in {'.md', '.markdown'}:
                        md_docs = self._load_markdown(file_path)
                        for doc in md_docs:
                            doc.metadata['gcs_path'] = gcs_path
                            doc.metadata['ingestion_metadata_id'] = metadata.id
                            doc.metadata['metadata_id'] = metadata.id  # Keep for backward compatibility
                            doc.metadata['document_id'] = document_id
                            doc.metadata['machine_model_ids'] = machine_model_ids
                            doc.metadata['machine_model_names'] = machine_model_names
                            if machine_model_names:
                                doc.metadata['machine_model'] = machine_model_names
                        documents.extend(md_docs)
                    
                except Exception as e:
                    logger.error(f"Error loading {metadata.filename} from database: {e}", exc_info=True)
                    continue
            
            logger.info(f"Loaded {len(documents)} document sections from database")
            
        finally:
            session.close()
        
        return documents
    
    def _load_from_gcs(self) -> List[Document]:
        """Load documents from GCS bucket."""
        from backend.utils.gcs_client import list_objects, download_to_file
        import tempfile
        
        logger.info(f"Loading documents from GCS bucket: {self.gcs_bucket}, prefix: {self.gcs_prefix}")
        
        # List all objects in GCS bucket with prefix
        object_infos = list_objects(self.gcs_bucket, self.gcs_prefix)
        object_names = [o.name for o in object_infos]
        
        # Filter to PDFs (case-insensitive)
        pdf_objects = [
            obj for obj in object_names
            if obj.lower().endswith('.pdf')
        ]
        
        logger.info(f"Found {len(pdf_objects)} PDF files in GCS bucket")
        
        documents = []
        temp_dir = tempfile.mkdtemp(prefix="ingest_gcs_")
        self.temp_files.append(temp_dir)  # Track for cleanup
        
        for obj_name in tqdm(pdf_objects, desc="Loading documents from GCS"):
            try:
                # Download to temporary file
                filename = os.path.basename(obj_name)
                temp_file_path = os.path.join(temp_dir, filename)
                
                gcs_uri = f"gs://{self.gcs_bucket}/{obj_name}"
                logger.debug(f"Downloading {gcs_uri} to {temp_file_path}")
                
                if not download_to_file(gcs_uri, temp_file_path):
                    logger.error(f"Failed to download {gcs_uri}")
                    continue
                
                self.temp_files.append(temp_file_path)  # Track for cleanup
                
                # Load document using existing logic
                file_path = Path(temp_file_path)
                file_ext = file_path.suffix.lower()
                
                if file_ext == '.pdf':
                    pdf_docs = SimpleDirectoryReader(input_files=[str(file_path)]).load_data()
                    for doc in pdf_docs:
                        doc.metadata['file_name'] = filename
                        doc.metadata['file_type'] = 'pdf'
                        doc.metadata['gcs_path'] = gcs_uri  # Store GCS path in metadata
                    documents.extend(pdf_docs)
                elif file_ext == '.docx' and DOCX_AVAILABLE:
                    docx_docs = self._load_docx(file_path)
                    for doc in docx_docs:
                        doc.metadata['gcs_path'] = gcs_uri
                    documents.extend(docx_docs)
                elif file_ext in {'.md', '.markdown'}:
                    md_docs = self._load_markdown(file_path)
                    for doc in md_docs:
                        doc.metadata['gcs_path'] = gcs_uri
                    documents.extend(md_docs)
                    
            except Exception as e:
                logger.error(f"Error loading {obj_name} from GCS: {e}", exc_info=True)
                continue
        
        logger.info(f"Loaded {len(documents)} document sections from GCS")
        return documents
    
    def _load_from_local(self) -> List[Document]:
        """Load documents from local directory (original behavior)."""
        documents = []
        
        # Get all supported files
        all_files = []
        for ext in self.supported_extensions:
            all_files.extend(list(self.data_dir.glob(f"**/*{ext}")))
        
        logger.info(f"Found {len(all_files)} files to process")
        
        for file_path in tqdm(all_files, desc="Loading documents"):
            try:
                file_ext = file_path.suffix.lower()
                file_name = file_path.name
                
                if file_ext == '.pdf':
                    # Use SimpleDirectoryReader for PDFs (existing logic)
                    pdf_docs = SimpleDirectoryReader(input_files=[str(file_path)]).load_data()
                    documents.extend(pdf_docs)
                    
                elif file_ext == '.docx' and DOCX_AVAILABLE:
                    docx_docs = self._load_docx(file_path)
                    documents.extend(docx_docs)
                    
                elif file_ext in {'.md', '.markdown'}:
                    md_docs = self._load_markdown(file_path)
                    documents.extend(md_docs)
                    
            except Exception as e:
                logger.error(f"Error loading {file_path}: {e}")
                continue
        
        return documents
    
    def cleanup_temp_files(self):
        """Clean up temporary files created during GCS downloads."""
        import shutil
        for temp_path in self.temp_files:
            try:
                if os.path.isdir(temp_path):
                    shutil.rmtree(temp_path)
                elif os.path.isfile(temp_path):
                    os.remove(temp_path)
            except Exception as e:
                logger.warning(f"Failed to cleanup temp file {temp_path}: {e}")
    
    def _load_docx(self, file_path: Path) -> List[Document]:
        """
        Load DOCX file and convert to Document objects with section/page metadata.
        """
        documents = []
        
        try:
            doc = DocxDocument(str(file_path))
            file_name = file_path.name
            
            # Extract text by paragraphs (for section tracking)
            paragraphs = []
            current_section = 1
            section_text = []
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                
                # Detect section headers (bold, larger font, or specific patterns)
                is_header = (
                    para.style.name.startswith('Heading') or
                    para.style.name.startswith('Title') or
                    para.runs and any(run.bold for run in para.runs) or
                    len(text) < 100 and text.isupper()
                )
                
                if is_header and section_text:
                    # Save previous section
                    section_text_combined = '\n'.join(section_text)
                    if section_text_combined.strip():
                        doc_obj = Document(
                            text=section_text_combined,
                            metadata={
                                'file_name': file_name,
                                'page_label': str(current_section),  # Use section number as page_label
                                'content_type': 'text',
                                'file_type': 'docx',
                                'section_number': current_section
                            }
                        )
                        documents.append(doc_obj)
                        current_section += 1
                        section_text = []
                
                section_text.append(text)
            
            # Add final section
            if section_text:
                section_text_combined = '\n'.join(section_text)
                if section_text_combined.strip():
                    doc_obj = Document(
                        text=section_text_combined,
                        metadata={
                            'file_name': file_name,
                            'page_label': str(current_section),
                            'content_type': 'text',
                            'file_type': 'docx',
                            'section_number': current_section
                        }
                    )
                    documents.append(doc_obj)
            
            # If no sections detected, create single document
            if not documents:
                full_text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
                if full_text.strip():
                    doc_obj = Document(
                        text=full_text,
                        metadata={
                            'file_name': file_name,
                            'page_label': '1',
                            'content_type': 'text',
                            'file_type': 'docx',
                            'section_number': 1
                        }
                    )
                    documents.append(doc_obj)
            
            logger.info(f"Loaded DOCX {file_name}: {len(documents)} sections")
            
        except Exception as e:
            logger.error(f"Error loading DOCX {file_path}: {e}")
        
        return documents
    
    def _load_markdown(self, file_path: Path) -> List[Document]:
        """
        Load Markdown file and convert to Document objects with section metadata.
        """
        documents = []
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            file_name = file_path.name
            
            # Split by markdown headers (# Header, ## Header, etc.)
            lines = content.split('\n')
            sections = []
            current_section = []
            current_section_num = 1
            
            for line in lines:
                # Check if line is a markdown header
                is_header = False
                header_level = 0
                
                if line.strip().startswith('#'):
                    # Count # symbols
                    header_level = len(line) - len(line.lstrip('#'))
                    if header_level <= 3:  # Only treat h1-h3 as section breaks
                        is_header = True
                
                if is_header and current_section:
                    # Save previous section
                    section_text = '\n'.join(current_section).strip()
                    if section_text:
                        sections.append({
                            'text': section_text,
                            'section_num': current_section_num,
                            'header': line.strip()
                        })
                        current_section_num += 1
                    current_section = []
                
                current_section.append(line)
            
            # Add final section
            if current_section:
                section_text = '\n'.join(current_section).strip()
                if section_text:
                    sections.append({
                        'text': section_text,
                        'section_num': current_section_num,
                        'header': ''
                    })
            
            # Create Document objects
            for section in sections:
                doc_obj = Document(
                    text=section['text'],
                    metadata={
                        'file_name': file_name,
                        'page_label': str(section['section_num']),  # Use section number as page_label
                        'content_type': 'text',
                        'file_type': 'markdown',
                        'section_number': section['section_num'],
                        'section_header': section['header']
                    }
                )
                documents.append(doc_obj)
            
            # If no sections detected, create single document
            if not documents and content.strip():
                doc_obj = Document(
                    text=content,
                    metadata={
                        'file_name': file_name,
                        'page_label': '1',
                        'content_type': 'text',
                        'file_type': 'markdown',
                        'section_number': 1
                    }
                )
                documents.append(doc_obj)
            
            logger.info(f"Loaded Markdown {file_name}: {len(documents)} sections")
            
        except Exception as e:
            logger.error(f"Error loading Markdown {file_path}: {e}")
        
        return documents


class TechnicalRAGPipeline:
    """High-performance RAG pipeline optimized for technical documentation with non-text content support."""
    
    def __init__(self, cache_dir="/root/.cache/huggingface/hub", config_path="config.yaml"):
        self.cache_dir = cache_dir
        self.embed_model = None
        self.embedding_model_name: str | None = None
        self.reranker = None
        self.index = None
        self.non_text_extractor = NonTextExtractor()
        self.text_preprocessor = TextPreprocessor()
        self.config = self._load_config(config_path)

        # Optional: when ingesting from a staging manifest, we populate this mapping so that
        # all non-text nodes (tables/images/captions) can include REQUIRED metadata fields.
        # Key: absolute local file path (string). Value: dict with document_id/machine_models/source_gcs.
        self._required_meta_by_local_path: dict[str, dict[str, Any]] = {}
        
        # Initialize Claude semantic rewriter (optional)
        claude_config = self.config.get("claude_rewriting", {})
        self.claude_rewriter = ClaudeSemanticRewriter(
            api_key=claude_config.get("api_key"),
            model=claude_config.get("model", "claude-3-5-sonnet-20241022"),
            enabled=claude_config.get("enabled", False),
            max_retries=claude_config.get("max_retries", 2),
            timeout=claude_config.get("timeout", 30)
        )
        
    def _load_config(self, config_path: str) -> Dict[str, Any]:
        """Load configuration from YAML file or use defaults."""
        default_config = {
            "qdrant": {
                "url": "http://localhost:6333",
                "collection_name": "technical_docs"
            },
            "models": {
                "embedding": "BAAI/bge-large-en-v1.5",
                "reranker": "BAAI/bge-reranker-large"
            },
            "chunking": {
                "chunk_size": 512,
                "chunk_overlap": 128
            }
        }
        
        if os.path.exists(config_path):
            try:
                with open(config_path, 'r') as f:
                    config = yaml.safe_load(f)
                return {**default_config, **config}
            except Exception as e:
                logger.warning(f"Failed to load config: {e}, using defaults")
        
        return default_config
        
    def initialize_models(self):
        """Initialize embedding and re-ranking models."""
        # Cloud Run safeguard: prevent ingestion on Cloud Run
        from backend.utils.cloud_run import is_cloud_run
        
        if is_cloud_run():
            raise RuntimeError("Ingestion cannot run on Cloud Run — must be executed on GPU externally.")
        
        logger.info("🚀 Initializing embedding model...")
        
        # Disable hf_transfer if not installed (RunPod issue)
        import shutil
        if os.environ.get('HF_HUB_ENABLE_HF_TRANSFER') == '1':
            logger.info("Disabling HF_HUB_ENABLE_HF_TRANSFER (package not installed)")
            os.environ['HF_HUB_ENABLE_HF_TRANSFER'] = '0'
        
        # Detect GPU with fallback to CPU if CUDA is incompatible
        import torch
        device = "cpu"  # Default to CPU
        try:
            if torch.cuda.is_available():
                # Test if CUDA actually works with a real operation
                try:
                    test_tensor = torch.zeros(1).cuda()
                    # Try a simple operation that requires kernel execution
                    result = test_tensor + 1
                    result.item()  # Force execution
                    del test_tensor, result
                    torch.cuda.empty_cache()
                    device = "cuda"
                    logger.info(f"🖥️ Using device: {device}")
                    logger.info(f"GPU: {torch.cuda.get_device_name(0)}")
                    logger.info(f"GPU Memory: {torch.cuda.get_device_properties(0).total_memory / 1e9:.2f} GB")
                except RuntimeError as cuda_error:
                    if "CUDA" in str(cuda_error) or "kernel" in str(cuda_error).lower():
                        logger.warning(f"⚠️ CUDA not compatible (kernel error), falling back to CPU: {cuda_error}")
                        device = "cpu"
                        logger.info(f"🖥️ Using device: {device}")
                    else:
                        raise
            else:
                logger.info(f"🖥️ Using device: {device} (CUDA not available)")
        except Exception as e:
            logger.warning(f"⚠️ CUDA test failed, falling back to CPU: {e}")
            device = "cpu"
            logger.info(f"🖥️ Using device: {device}")
        
        cache_path = os.path.expanduser(self.cache_dir)
        
        # Try multiple approaches
        model_options = [
            ("BAAI/bge-large-en-v1.5", "BGE Large"),
            ("BAAI/bge-base-en-v1.5", "BGE Base"),
            ("all-MiniLM-L6-v2", "MiniLM"),
            ("all-mpnet-base-v2", "MPNet")
        ]
        
        for model_name, display_name in model_options:
            try:
                logger.info(f"Trying model: {display_name} ({model_name})")
                
                # Method 1: Direct load without sentence-transformers prefix
                try:
                    # Force CPU if CUDA is incompatible
                    if device == "cuda":
                        try:
                            # Test CUDA compatibility
                            import torch
                            test = torch.zeros(1).cuda()
                            del test
                        except Exception:
                            logger.warning(f"CUDA incompatible, forcing CPU for {display_name}")
                            device = "cpu"
                    
                    self.embed_model = HuggingFaceEmbedding(
                        model_name=model_name,
                        cache_folder=self.cache_dir,
                        trust_remote_code=True,
                        device=device
                    )
                    self.embedding_model_name = model_name
                    logger.info(f"✅ Successfully loaded: {display_name} on {device}")
                    break
                except Exception as e1:
                    logger.debug(f"Method 1 failed: {e1}")
                    
                    # Method 2: Try with full sentence-transformers path
                    if not model_name.startswith("sentence-transformers/"):
                        try:
                            full_name = f"sentence-transformers/{model_name}"
                            self.embed_model = HuggingFaceEmbedding(
                                model_name=full_name,
                                cache_folder=self.cache_dir,
                                trust_remote_code=True,
                                device=device
                            )
                            self.embedding_model_name = full_name
                            logger.info(f"✅ Successfully loaded: {display_name} on {device}")
                            break
                        except Exception as e2:
                            logger.debug(f"Method 2 failed: {e2}")
                            raise e1
                    else:
                        raise e1
                        
            except Exception as e:
                logger.warning(f"Failed to load {display_name}: {str(e)[:100]}")
                continue
        
        if not self.embed_model:
            logger.error("All model loading attempts failed. Trying emergency fallback...")
            # Emergency fallback - use any available model
            try:
                self.embed_model = HuggingFaceEmbedding(model_name="all-MiniLM-L6-v2")
                self.embedding_model_name = "all-MiniLM-L6-v2"
                logger.info("✅ Loaded with emergency fallback")
            except:
                raise RuntimeError("Could not load any embedding model. Check internet connection and HuggingFace access.")
        
        # Try to initialize re-ranker (optional)
        try:
            logger.info("🎯 Initializing re-ranker...")
            reranker_model = self.config.get("models", {}).get("reranker", "BAAI/bge-reranker-large")
            # Use CPU if CUDA was incompatible
            reranker_device = device if device == "cpu" else device
            try:
                try:
                    # Newer sentence-transformers CrossEncoder may not support cache_folder
                    self.reranker = CrossEncoder(
                        reranker_model,
                        cache_folder=self.cache_dir,
                        device=reranker_device
                    )
                except TypeError:
                    self.reranker = CrossEncoder(
                        reranker_model,
                        device=reranker_device
                    )
                logger.info(f"✅ Re-ranker loaded successfully on {reranker_device}")
            except RuntimeError as cuda_error:
                if "CUDA" in str(cuda_error) or "cuda" in str(cuda_error).lower():
                    logger.warning(f"⚠️ CUDA incompatible for reranker, using CPU: {cuda_error}")
                    try:
                        self.reranker = CrossEncoder(
                            reranker_model,
                            cache_folder=self.cache_dir,
                            device="cpu"
                        )
                    except TypeError:
                        self.reranker = CrossEncoder(
                            reranker_model,
                            device="cpu"
                        )
                    logger.info(f"✅ Re-ranker loaded successfully on CPU")
                else:
                    raise
        except Exception as e:
            logger.warning(f"Re-ranker not available: {e}")
            self.reranker = None
        
        # Set global embedding model
        Settings.embed_model = self.embed_model
        logger.info("✅ Models initialized successfully")
    
    def process_non_text_content(self, data_dir: str) -> Tuple[List[Dict], List[Dict], List[Dict]]:
        """Process non-text content (tables, images, captions) from documents."""
        logger.info("📊 Processing non-text content...")
        
        all_tables = []
        all_images = []
        all_captions = []
        
        # Find all PDF files (recursive; GCS staging uses nested directories like documents/<metadata_id>/<file>.pdf)
        pdf_files = list(Path(data_dir).rglob("*.pdf"))
        logger.info(f"Found {len(pdf_files)} PDF files to process")
        
        for pdf_path in pdf_files:
            logger.info(f"Processing {pdf_path.name}...")
            
            try:
                # Extract tables
                tables = self.non_text_extractor.extract_tables_from_pdf(str(pdf_path))
                all_tables.extend(tables)
                logger.info(f"Extracted {len(tables)} tables from {pdf_path.name}")
                
                # Extract images
                images = self.non_text_extractor.extract_images_from_pdf(str(pdf_path))
                all_images.extend(images)
                logger.info(f"Extracted {len(images)} images from {pdf_path.name}")
                
                # Extract captions
                captions = self.non_text_extractor.extract_figure_captions(str(pdf_path))
                all_captions.extend(captions)
                logger.info(f"Extracted {len(captions)} captions from {pdf_path.name}")
                
            except Exception as e:
                logger.error(f"Failed to process {pdf_path.name}: {e}")
                continue
        
        logger.info(f"✅ Non-text processing complete: {len(all_tables)} tables, {len(all_images)} images, {len(all_captions)} captions")
        return all_tables, all_images, all_captions
    
    def create_non_text_nodes(self, tables: List[Dict], images: List[Dict], captions: List[Dict]) -> List[TextNode]:
        """Create TextNode objects for non-text content to be embedded."""
        nodes = []

        def _required_meta_for_source_path(source_path: str) -> dict[str, Any]:
            try:
                key = str(Path(source_path).resolve())
            except Exception:
                key = source_path
            meta = self._required_meta_by_local_path.get(key)
            if meta:
                return meta
            # Loud fallback: stable UUID5 derived from the local source_path string.
            stable_id = _stable_uuid5_from_string(str(source_path))
            logger.warning(
                "Non-text node missing required metadata mapping; using deterministic fallback",
                extra={"source_path": source_path, "fallback_document_id": stable_id},
            )
            return {
                "document_id": stable_id,
                "machine_model_ids": [],
                "machine_model_names": [],
                "machine_models": [],
                "machine_model": [],
                "source_gcs": None,
                "gcs_path": None,
            }
        
        # Process tables
        for table in tables:
            # Create text representation of table
            table_text = f"Table from {Path(table['source_path']).name}, page {table['page_number']}:\n{table['table_markdown']}"
            
            node = TextNode(
                text=table_text,
                metadata={
                    "content_type": "table",
                    "source_path": table["source_path"],
                    "page_number": table["page_number"],
                    "table_index": table["table_index"],
                    "row_count": table["row_count"],
                    "column_count": table["column_count"],
                    "table_json": table["table_json"],
                    **_required_meta_for_source_path(table.get("source_path")),
                }
            )
            nodes.append(node)
        
        # Process figure captions
        for caption in captions:
            # Clean caption text to remove boilerplate
            caption_text = self.text_preprocessor.clean_text(caption["caption_text"])
            if not self.text_preprocessor.should_skip_node(caption_text):
                node = TextNode(
                    text=caption_text,
                    metadata={
                        "content_type": "figure_caption",
                        "source_path": caption["source_path"],
                        "page_number": caption["page_number"],
                        "line_number": caption["line_number"],
                        **_required_meta_for_source_path(caption.get("source_path")),
                    }
                )
                nodes.append(node)
        
        # Process images (create text nodes for captions and metadata only - no base64)
        for image in images:
            image_text = f"Image from {Path(image['source_path']).name}, page {image['page_number']}: {image['caption']}"
            
            node = TextNode(
                text=image_text,
                metadata={
                    "content_type": "image",
                    "source_path": image["source_path"],
                    "page_number": image["page_number"],
                    "image_index": image["image_index"],
                    "width": image["width"],
                    "height": image["height"],
                    "saved_path": image.get("saved_path"),
                    "bbox": str(image.get("bbox")) if image.get("bbox") else None,
                    **_required_meta_for_source_path(image.get("source_path")),
                }
            )
            nodes.append(node)
        
        return nodes
    
    def setup_qdrant_storage(self) -> StorageContext:
        """Setup Qdrant vector store for hybrid search."""
        try:
            # Initialize Qdrant client
            qdrant_url = self.config["qdrant"]["url"]
            collection_name = self.config["qdrant"]["collection_name"]
            
            client = qdrant_client.QdrantClient(url=qdrant_url)
            
            # Create vector store
            vector_store = QdrantVectorStore(
                client=client,
                collection_name=collection_name
            )
            
            # Create storage context
            storage_context = StorageContext.from_defaults(
                vector_store=vector_store,
                docstore=SimpleDocumentStore(),
                index_store=SimpleIndexStore()
            )
            
            logger.info(f"✅ Qdrant storage configured: {qdrant_url}/{collection_name}")
            return storage_context
            
        except Exception as e:
            logger.warning(f"Qdrant not available, using local storage: {e}")
            return None
    
    def build_index(self, data_dir="data", storage_dir="latest_model", use_qdrant=False):
        """Build or load vector index with optimized chunking and non-text content."""

        # Make relative paths resilient to current working directory
        if data_dir == "data":
            data_dir = DEFAULT_DATA_DIR
        if storage_dir == "latest_model":
            storage_dir = DEFAULT_STORAGE_DIR
        
        # Initialize models
        self.initialize_models()
        
        # Setup storage context
        storage_context = None
        if use_qdrant:
            storage_context = self.setup_qdrant_storage()
        
        # For local storage: always rebuild (clear old index first)
        if not use_qdrant and os.path.exists(storage_dir):
            logger.info(f"🗑️  Clearing old index from {storage_dir} for fresh rebuild...")
            shutil.rmtree(storage_dir)
            logger.info("✅ Old index cleared - ready for fresh build")
        
        # Create storage directory if it doesn't exist
        if not use_qdrant:
            os.makedirs(storage_dir, exist_ok=True)
        
        print("\n" + "="*70)
        print("📥 BUILDING NEW RAG INDEX")
        print("="*70)
        
        # Step 1: Load Documents (PDF, DOCX, Markdown)
        print("\n[Step 1/7] 📄 Loading documents (PDF, DOCX, Markdown)...")
        
        # Optional: deterministic ingestion input via staging manifest
        manifest_path = os.getenv("INGEST_DOC_MANIFEST_PATH")
        if manifest_path and os.path.exists(manifest_path):
            print(f"   📋 Loading documents from staging manifest: {manifest_path}")
            loader = DocumentLoader(data_dir=data_dir, manifest_path=manifest_path)
            documents = loader.load_documents(use_database=False)
            use_gcs = False
            # Populate non-text metadata mapping for tables/images/captions
            try:
                with open(manifest_path, "r", encoding="utf-8") as f:
                    manifest = json.load(f) or {}
                entries = manifest.get("documents", []) if isinstance(manifest, dict) else []
                mapping: dict[str, dict[str, Any]] = {}
                for entry in entries if isinstance(entries, list) else []:
                    try:
                        lp = entry.get("local_path")
                        if not lp:
                            continue
                        mm_names = entry.get("machine_model_names") or entry.get("machine_models") or []
                        mm_ids = entry.get("machine_model_ids") or []
                        if not isinstance(mm_names, list):
                            mm_names = []
                        if not isinstance(mm_ids, list):
                            mm_ids = []
                        mapping[str(Path(lp).resolve())] = {
                            "document_id": str(entry.get("document_id")) if entry.get("document_id") is not None else "0",
                            "machine_model_ids": [str(x) for x in mm_ids if isinstance(x, str) and x.strip()],
                            "machine_model_names": [str(x) for x in mm_names if isinstance(x, str) and x.strip()],
                            "machine_models": [str(x) for x in mm_names if isinstance(x, str) and x.strip()],
                            "source_gcs": entry.get("source_gcs"),
                            "gcs_path": entry.get("source_gcs"),
                            # Orchestrator filter uses machine_model as list[str]
                            "machine_model": [str(x) for x in mm_names if isinstance(x, str) and x.strip()],
                        }
                    except Exception:
                        continue
                self._required_meta_by_local_path = mapping
            except Exception as e:
                logger.warning(f"Failed to load manifest for non-text metadata propagation: {e}")
        else:
            # Check if GCS is configured for document storage
            from backend.config.env import settings
            use_gcs = bool(settings.DOCS_GCS_BUCKET)

            # Always try to load from database first (only processes documents in DB)
            # Falls back to GCS/local if database loading fails or returns no documents
            loader = None
            if use_gcs:
                print(f"   📦 Loading from database (documents with GCS paths)...")
                loader = DocumentLoader(
                    gcs_bucket=settings.DOCS_GCS_BUCKET,
                    gcs_prefix=settings.DOCS_GCS_PREFIX
                )
            else:
                print(f"   📁 Loading from database (documents with GCS paths)...")
                loader = DocumentLoader(data_dir=data_dir)

            # Load from database (preferred) - only processes documents in database
            # This ensures we only ingest documents that are tracked in the database
            documents = loader.load_documents(use_database=True)

        # Safety: never build/upload an empty index unless explicitly allowed
        allow_empty = os.getenv("ALLOW_EMPTY_INDEX", "false").lower() in {"1", "true", "yes", "on"}
        if len(documents) == 0 and not allow_empty:
            raise RuntimeError(
                "No documents loaded (0). Refusing to build/upload an empty index. "
                "Fix by either: (a) run from repo root so data/ is found, "
                "(b) set DOCS_GCS_BUCKET/DOCS_GCS_PREFIX to load from GCS, "
                "and/or (c) run Cloud SQL Auth Proxy + set a TCP DATABASE_URL if loading from DB. "
                "If you truly want an empty index, set ALLOW_EMPTY_INDEX=true."
            )
        
        # Note: No need to sync to database since we're loading FROM database
        # Documents are already tracked in DocumentIngestionMetadata and Document tables
        
        # Count by file type
        file_types = {}
        for doc in documents:
            file_type = doc.metadata.get('file_type', 'pdf')
            file_types[file_type] = file_types.get(file_type, 0) + 1
        
        type_summary = ", ".join([f"{count} {ftype.upper()}" for ftype, count in file_types.items()])
        print(f"   ✅ Loaded {len(documents)} document sections ({type_summary})")
        logger.info(f"Loaded {len(documents)} documents: {type_summary}")
        
        # Cleanup temp files if loaded from GCS
        if use_gcs:
            try:
                loader.cleanup_temp_files()
            except Exception as e:
                logger.warning(f"Failed to cleanup temp files: {e}")
        
        # Step 2: Enhanced AI-Powered Preprocessing
        print("\n[Step 2/7] 🧹 Enhanced preprocessing (TOC removal, artifact fixing, normalization)...")
        preprocessed_docs = []
        skipped_pages = 0
        skip_reasons = {}
        
        for doc in documents:
            original_text = doc.text or ""
            
            # Enhanced cleaning with metadata for context-aware processing
            cleaned_text = self.text_preprocessor.clean_text(original_text, metadata=doc.metadata)
            
            # Check if page should be skipped (with reason tracking)
            if self.text_preprocessor.is_low_content_page(cleaned_text):
                skip_reasons['low_content'] = skip_reasons.get('low_content', 0) + 1
                skipped_pages += 1
                logger.debug(f"Skipping low-content page: {doc.metadata.get('file_name', 'unknown')}")
                continue
            
            # Check for first page without content
            if self.text_preprocessor.is_first_page_without_content(cleaned_text, metadata=doc.metadata):
                skip_reasons['first_page_no_content'] = skip_reasons.get('first_page_no_content', 0) + 1
                skipped_pages += 1
                logger.debug(f"Skipping first page without content: {doc.metadata.get('file_name', 'unknown')}")
                continue
            
            # Create new document with cleaned text
            if cleaned_text:  # Only add if there's content left after cleaning
                new_doc = Document(
                    text=cleaned_text,
                    metadata=doc.metadata
                )
                preprocessed_docs.append(new_doc)
        
        skip_summary = ", ".join([f"{reason}: {count}" for reason, count in skip_reasons.items()]) if skip_reasons else "none"
        print(f"   ✅ Preprocessed {len(preprocessed_docs)} documents ({skipped_pages} pages skipped: {skip_summary})")
        logger.info(f"Preprocessed {len(preprocessed_docs)} documents, skipped {skipped_pages} pages ({skip_summary})")
        
        # Step 3: Extract Non-Text Content
        print("\n[Step 3/7] 🖼️  Extracting tables, images, and captions...")
        print("   This may take a few minutes...")
        tables, images, captions = self.process_non_text_content(data_dir)
        print(f"   ✅ Extracted {len(tables)} tables, {len(images)} images, {len(captions)} captions")
        
        # Step 4: Create Non-Text Nodes
        print("\n[Step 4/7] 📊 Creating searchable nodes from extracted content...")
        non_text_nodes = self.create_non_text_nodes(tables, images, captions)
        print(f"   ✅ Created {len(non_text_nodes)} non-text nodes")
        logger.info(f"Created {len(non_text_nodes)} non-text nodes")
        
        # Step 5: Smart Chunking with Text Nodes
        print("\n[Step 5/7] 🧠 Smart chunking and filtering...")
        chunk_size = self.config.get("chunking", {}).get("chunk_size", 350)
        chunk_overlap = self.config.get("chunking", {}).get("chunk_overlap", 88)
        
        smart_splitter = SmartChunkSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            preprocessor=self.text_preprocessor
        )
        
        print(f"   - Chunk size: {chunk_size} characters")
        print(f"   - Chunk overlap: {chunk_overlap} characters")
        print(f"   - Preserving structured content (tables, code blocks, numbered steps)...")
        
        # Get nodes from documents using smart chunking
        text_nodes = smart_splitter.get_nodes_from_documents(preprocessed_docs, show_progress=True)
        
        # Filter out short/low-quality nodes (only for text nodes, not tables/images)
        filtered_nodes = []
        skipped_nodes = 0
        skip_reasons_chunks = {}
        
        for node in text_nodes:
            # Skip filtering for non-text content types (tables, images, captions are already handled separately)
            content_type = node.metadata.get("content_type", "text")
            if content_type != "text":
                filtered_nodes.append(node)  # Don't filter non-text content
            else:
                # Enhanced skip check with reason tracking
                should_skip, skip_reason = self.text_preprocessor.should_skip_node(
                    node.text, 
                    metadata=node.metadata
                )
                if should_skip:
                    skipped_nodes += 1
                    skip_reasons_chunks[skip_reason] = skip_reasons_chunks.get(skip_reason, 0) + 1
                    logger.debug(f"Skipping node from {node.metadata.get('file_name', 'unknown')}: {skip_reason}")
                else:
                    filtered_nodes.append(node)
        
        skip_summary_chunks = ", ".join([f"{reason}: {count}" for reason, count in skip_reasons_chunks.items()]) if skip_reasons_chunks else "none"
        print(f"   ✅ Created {len(filtered_nodes)} text nodes ({skipped_nodes} filtered: {skip_summary_chunks})")
        logger.info(f"Created {len(filtered_nodes)} text nodes, filtered {skipped_nodes} nodes ({skip_summary_chunks})")
        
        # Step 5.5: Optional Claude Semantic Rewriting
        if self.claude_rewriter.enabled:
            print("\n[Step 5.5/7] 🤖 Claude semantic rewriting (improving clarity while preserving meaning)...")
            print("   - This step uses Claude API to enhance text clarity")
            print("   - Structured content (tables, code, lists) will be preserved as-is")
            print("   - Estimated time: 1-3 minutes per 100 chunks")
            
            rewritten_nodes, rewrite_stats = self.claude_rewriter.rewrite_nodes(filtered_nodes, show_progress=True)
            
            print(f"   ✅ Rewriting complete:")
            print(f"      - Rewritten: {rewrite_stats['rewritten']} chunks")
            print(f"      - Preserved (structured): {rewrite_stats['structured']} chunks")
            print(f"      - Skipped: {rewrite_stats['skipped']} chunks")
            print(f"      - Failed (using original): {rewrite_stats['failed']} chunks")
            
            # Use rewritten nodes for embedding
            filtered_nodes = rewritten_nodes
        else:
            logger.debug("Claude rewriting disabled, skipping rewrite step")
        
        # Step 6: Create Vector Embeddings (LONGEST STEP)
        print("\n[Step 6/7] 🧠 Generating embeddings and building vector index...")
        print(f"   - Processing {len(filtered_nodes)} text nodes + {len(non_text_nodes)} non-text nodes...")
        print(f"   - This is the LONGEST step (embedding generation)")
        print(f"   - Expected time: 5-15 minutes on GPU, 30-60 minutes on CPU")
        print(f"   - Watch for progress below...")
        print("")
        
        # Record start time
        import time
        start_time = time.time()
        
        # Combine all nodes
        all_nodes = filtered_nodes + non_text_nodes
        
        # Create index from nodes (LlamaIndex API)
        if storage_context:
            # Create index with storage context
            self.index = VectorStoreIndex(
                nodes=[],
                storage_context=storage_context,
                show_progress=True
            )
        else:
            # Create index without storage context (will use default)
            self.index = VectorStoreIndex(
                nodes=[],
                show_progress=True
            )
        
        # Insert all nodes into the index (batch insert for better performance)
        print(f"   Inserting {len(all_nodes)} nodes into index...")
        batch_size = 100  # Insert in batches
        for i in tqdm(range(0, len(all_nodes), batch_size), desc="   Inserting nodes", unit="batch"):
            batch = all_nodes[i:i + batch_size]
            try:
                self.index.insert_nodes(batch)
            except RuntimeError as e:
                if "CUDA" in str(e) or "cuda" in str(e).lower():
                    logger.error(f"CUDA error during embedding - device should have been set to CPU")
                    logger.error(f"Error: {e}")
                    raise RuntimeError("CUDA incompatible. The embedding model should use CPU. Check device detection.") from e
                else:
                    raise
        
        elapsed = time.time() - start_time
        print(f"\n   ✅ Vector index created in {elapsed:.1f} seconds ({elapsed/60:.1f} minutes)")
        print(f"   ⚡ Processing speed: {len(all_nodes) / elapsed:.2f} nodes/sec")
        
        # Persist the index (only for local storage)
        print(f"\n💾 Saving index to disk...")
        if not use_qdrant:
            self.index.storage_context.persist(persist_dir=storage_dir)
            print(f"   ✅ Index saved to: {storage_dir}")
            logger.info("✅ Index created and saved locally")

            # Write a build manifest for verification + promotion workflows
            try:
                from datetime import datetime, timezone
                build_ts = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H-%M-%SZ")

                # Best-effort document count
                docs_count = None
                manifest_path = os.getenv("INGEST_DOC_MANIFEST_PATH")
                if manifest_path and os.path.exists(manifest_path):
                    with open(manifest_path, "r", encoding="utf-8") as f:
                        m = json.load(f) or {}
                    entries = m.get("documents", []) if isinstance(m, dict) else []
                    if isinstance(entries, list):
                        docs_count = len(entries)
                if docs_count is None:
                    docs_count = len({d.metadata.get("file_name") for d in documents if hasattr(d, "metadata") and d.metadata})

                chunk_count = len(all_nodes) if "all_nodes" in locals() else None

                manifest = {
                    "build_timestamp": build_ts,
                    "embedding_model": self.embedding_model_name or self.config.get("models", {}).get("embedding"),
                    "num_documents": docs_count,
                    "num_chunks": chunk_count,
                }
                manifest_path_out = Path(storage_dir) / "index_manifest.json"
                with open(manifest_path_out, "w", encoding="utf-8") as f:
                    json.dump(manifest, f, indent=2, sort_keys=True)
                logger.info(f"Wrote index manifest to {manifest_path_out}")
            except Exception as e:
                logger.warning(f"Failed to write index_manifest.json: {e}", exc_info=True)
        else:
            print(f"   ✅ Index saved to: Qdrant")
            logger.info("✅ Index created and saved to Qdrant")
        
        # --------------------------
        # STEP 7 — METADATA UPDATE
        # --------------------------
        metadata_updated = 0
        metadata_needs_review = 0
        
        if os.environ.get("DISABLE_METADATA_UPDATE", "1") == "1":
            logger.warning("Skipping metadata update step (DISABLE_METADATA_UPDATE=1).")
            print("\n[Step 7/7] 📝 Updating document metadata...")
            print("   ⚠️  Metadata update disabled (DISABLE_METADATA_UPDATE=1)")
        else:
            # Lazy import — only if needed
            from backend.utils.document_metadata import ensure_metadata_entry
            
            print("\n[Step 7/7] 📝 Updating document metadata...")
            
            # Collect all unique filenames from documents
            unique_filenames = set()
            for doc in documents:
                filename = doc.metadata.get('file_name')
                if filename:
                    unique_filenames.add(filename)
            
            # Also check from nodes (in case some documents were filtered out)
            for node in all_nodes:
                filename = node.metadata.get('file_name')
                if filename:
                    unique_filenames.add(filename)
            
            # Update metadata for each document
            for filename in unique_filenames:
                try:
                    meta_entry = ensure_metadata_entry(filename)
                    metadata_updated += 1
                    if meta_entry.get("requires_admin_review"):
                        metadata_needs_review += 1
                except Exception as e:
                    logger.warning(f"Failed to update metadata for {filename}: {e}")
            
            print(f"   ✅ Updated metadata for {metadata_updated} documents")
            if metadata_needs_review > 0:
                print(f"   ⚠️  {metadata_needs_review} documents require admin review (missing machine_model)")
        
        # Final summary
        total_time = time.time() - start_time
        print("\n" + "="*70)
        print("✅ INGESTION COMPLETE!")
        print("="*70)
        print(f"📊 Documents loaded: {len(documents)}")
        print(f"📊 Documents after preprocessing: {len(preprocessed_docs)} ({skipped_pages} low-content pages skipped)")
        print(f"📊 Text nodes created: {len(filtered_nodes)} ({skipped_nodes} filtered)")
        print(f"📊 Non-text nodes: {len(non_text_nodes)}")
        print(f"📊 Total nodes indexed: {len(all_nodes)}")
        if os.environ.get("DISABLE_METADATA_UPDATE", "1") == "1":
            print(f"📝 Metadata update: Skipped (DISABLE_METADATA_UPDATE=1)")
        else:
            print(f"📝 Metadata updated: {metadata_updated} documents")
            if metadata_needs_review > 0:
                print(f"⚠️  Documents needing review: {metadata_needs_review}")
        if self.claude_rewriter.enabled:
            print(f"🤖 Claude rewriting: Enabled (check logs for rewrite statistics)")
        print(f"⏱️  Total time: {total_time:.1f} seconds ({total_time/60:.1f} minutes)")
        if not use_qdrant:
            print(f"📁 Storage location: {storage_dir}")
            print(f"\n💡 Two-Pod Workflow:")
            print(f"   1. git add {storage_dir}/")
            print(f"   2. git commit -m 'Update RAG index'")
            print(f"   3. git push")
            print(f"   4. Switch to cheap pod → git pull → streamlit run app.py")
        print(f"📂 Extracted content: extracted_content/")
        print(f"🔍 Ready to query!")
        print("="*70 + "\n")
        
        return self.index
    
    def hybrid_search(self, query: str, top_k: int = 10, content_types: List[str] = None) -> List[NodeWithScore]:
        """Perform hybrid search across text, tables, and images."""
        if not self.index:
            raise RuntimeError("Index not built. Call build_index() first.")
        
        # Default to search all content types
        if content_types is None:
            content_types = ["text", "table", "image", "figure_caption"]
        
        # Perform vector search
        retriever = self.index.as_retriever(similarity_top_k=top_k * 2)  # Get more for re-ranking
        nodes = retriever.retrieve(query)
        
        # Filter by content type if specified
        if content_types:
            filtered_nodes = []
            for node in nodes:
                content_type = node.metadata.get("content_type", "text")
                if content_type in content_types or content_type == "text":
                    filtered_nodes.append(node)
            nodes = filtered_nodes[:top_k]
        
        # Apply re-ranking if available
        if self.reranker and len(nodes) > 1:
            logger.info("🎯 Applying re-ranking...")
            try:
                # Prepare query-document pairs for re-ranking
                pairs = [(query, node.text) for node in nodes]
                scores = self.reranker.predict(pairs)
                
                # Sort by re-ranking scores
                scored_nodes = list(zip(nodes, scores))
                scored_nodes.sort(key=lambda x: x[1], reverse=True)
                nodes = [node for node, score in scored_nodes[:top_k]]
                
            except Exception as e:
                logger.warning(f"Re-ranking failed: {e}")
        
        return nodes[:top_k]
    
    def backup_storage(self, storage_dir: str, backup_name: str = None):
        """Create a backup of the storage directory."""
        if not os.path.exists(storage_dir):
            logger.warning(f"Storage directory {storage_dir} does not exist")
            return None
            
        if backup_name is None:
            backup_name = f"rag_backup_{int(time.time())}"
        
        backup_path = f"/workspace/{backup_name}.tar.gz"
        
        try:
            with tarfile.open(backup_path, "w:gz") as tar:
                tar.add(storage_dir, arcname=os.path.basename(storage_dir))
            
            logger.info(f"✅ Backup created: {backup_path}")
            return backup_path
        except Exception as e:
            logger.error(f"Failed to create backup: {e}")
            return None
    
    def restore_storage(self, backup_path: str, storage_dir: str):
        """Restore storage from backup."""
        try:
            with tarfile.open(backup_path, "r:gz") as tar:
                tar.extractall(os.path.dirname(storage_dir))
            logger.info(f"✅ Storage restored from {backup_path}")
            return True
        except Exception as e:
            logger.error(f"Failed to restore backup: {e}")
            return False


def _sync_gcs_documents_to_db(documents: List[Document], loader: DocumentLoader):
    """
    Sync GCS documents to database.
    Creates/updates Document and DocumentIngestionMetadata records.
    """
    from backend.utils.db import SessionLocal, Document as DBDocument, DocumentIngestionMetadata
    from backend.utils.gcs_client import parse_gcs_path
    import uuid
    from datetime import datetime
    
    session = SessionLocal()
    try:
        # Group documents by filename (from GCS path)
        documents_by_filename = {}
        for doc in documents:
            gcs_path = doc.metadata.get('gcs_path')
            if not gcs_path:
                continue
            
            filename = doc.metadata.get('file_name')
            if not filename:
                # Extract filename from GCS path
                _, blob_name = parse_gcs_path(gcs_path)
                filename = os.path.basename(blob_name) if blob_name else None
            
            if filename:
                if filename not in documents_by_filename:
                    documents_by_filename[filename] = {
                        'gcs_path': gcs_path,
                        'documents': []
                    }
                documents_by_filename[filename]['documents'].append(doc)
        
        logger.info(f"Syncing {len(documents_by_filename)} documents to database...")
        
        for filename, doc_info in documents_by_filename.items():
            try:
                gcs_path = doc_info['gcs_path']
                
                # Parse document_id from GCS path if it matches convention: {prefix}{document_id}/{filename}
                # Example: documents/abc-123-def/filename.pdf -> document_id = abc-123-def
                document_id = None
                bucket_name, blob_name = parse_gcs_path(gcs_path)
                if blob_name:
                    parts = blob_name.split('/')
                    if len(parts) >= 2:
                        # Check if first part looks like a UUID/metadata_id
                        potential_id = parts[0]
                        if len(potential_id) > 10:  # Likely a UUID or similar ID
                            document_id = potential_id
                
                # Check if Document exists
                db_doc = session.query(DBDocument).filter(DBDocument.file_name == filename).first()
                
                if not db_doc:
                    # Create new Document record
                    db_doc = DBDocument(
                        file_name=filename,
                        gcs_path=gcs_path,
                        display_name=filename,
                        is_active=True,
                        requires_admin_review=False,
                    )
                    session.add(db_doc)
                    session.flush()  # Get the ID
                    logger.info(f"Created Document record for {filename}")
                else:
                    # Update GCS path if not set
                    if not db_doc.gcs_path:
                        db_doc.gcs_path = gcs_path
                        logger.info(f"Updated Document record with GCS path for {filename}")
                
                # Ensure DocumentIngestionMetadata exists
                if document_id:
                    # Try to find by ID
                    metadata = session.query(DocumentIngestionMetadata).filter(
                        DocumentIngestionMetadata.id == document_id
                    ).first()
                else:
                    # Try to find by filename
                    metadata = session.query(DocumentIngestionMetadata).filter(
                        DocumentIngestionMetadata.filename == filename
                    ).first()
                
                if not metadata:
                    # Create new metadata record
                    metadata_id = document_id or str(uuid.uuid4())
                    metadata = DocumentIngestionMetadata(
                        id=metadata_id,
                        filename=filename,
                        machine_model="GENERAL",  # Default, can be updated later
                        status="COMPLETE",  # Assume complete if in GCS
                        file_size_bytes=None,  # Could extract from GCS blob if needed
                    )
                    session.add(metadata)
                    logger.info(f"Created DocumentIngestionMetadata for {filename}")
                else:
                    # Update status if needed
                    if metadata.status not in ("COMPLETE", "READY"):
                        metadata.status = "COMPLETE"
                    logger.debug(f"DocumentIngestionMetadata already exists for {filename}")
                
            except Exception as e:
                logger.error(f"Failed to sync {filename} to database: {e}", exc_info=True)
                session.rollback()
                continue
        
        session.commit()
        logger.info(f"✅ Synced {len(documents_by_filename)} documents to database")
        
    finally:
        session.close()


def _env_bool(name: str, default: bool = False) -> bool:
    v = os.getenv(name)
    if v is None:
        return default
    return v.strip().lower() in {"1", "true", "yes", "on"}


def _normalize_prefix(prefix: str) -> str:
    p = (prefix or "").strip()
    if not p:
        return ""
    p = p.lstrip("/")
    return p if p.endswith("/") else f"{p}/"


def _stable_uuid5_from_string(s: str) -> str:
    """
    Deterministic, stable UUID string derived from an input string.

    Used as a fallback when DB doesn't provide a canonical document_id.
    """
    import uuid
    return str(uuid.uuid5(uuid.NAMESPACE_URL, s))


def _safe_rel_from_gcs_object(docs_prefix: str, object_name: str) -> str:
    """
    Convert a GCS object name into a safe relative path under docs_prefix.

    Prevents path traversal / absolute paths from being interpreted as filesystem paths.
    If unsafe, falls back to a deterministic hashed filename under "__unsafe__/".
    """
    from pathlib import PurePosixPath
    import hashlib

    rel = object_name
    if docs_prefix and rel.startswith(docs_prefix):
        rel = rel[len(docs_prefix):].lstrip("/")

    p = PurePosixPath(rel)
    parts = p.parts
    if not parts or p.is_absolute() or any(seg in {"..", ""} for seg in parts):
        ext = PurePosixPath(object_name).suffix.lower()
        h = hashlib.sha1(object_name.encode("utf-8")).hexdigest()[:16]
        return f"__unsafe__/{h}{ext}"
    return p.as_posix()


def _parse_machine_models(raw: Any) -> list[str]:
    if raw is None:
        return []
    if isinstance(raw, list):
        return [m for m in raw if isinstance(m, str) and m.strip()]
    if isinstance(raw, str):
        r = raw.strip()
        if not r:
            return []
        try:
            parsed = json.loads(r)
            if isinstance(parsed, list):
                return [m for m in parsed if isinstance(m, str) and m.strip()]
        except Exception:
            pass
        # Comma-separated fallback
        return [m.strip() for m in r.split(",") if m.strip()]
    return []


def _resolve_authoritative_doc_metadata(
    session,
    docs_bucket: str,
    docs_prefix: str,
    object_name: str,
    object_custom_metadata: dict[str, Any] | None,
) -> dict[str, Any]:
    """
    Resolve document_id (non-empty string) and machine model metadata for a GCS object.

    Source of truth order:
    1) DB Document (match by gcs_path)
    2) DB DocumentIngestionMetadata (match by metadata_id parsed from object path, or file_path)
    3) GCS object custom metadata (document_id, machine_model_ids, machine_model_names / machine_models)
    4) Deterministic fallback (UUID5(gs://bucket/object); machine_models=[])
    """
    source_gcs = f"gs://{docs_bucket}/{object_name}"

    db_doc_id: str | None = None
    machine_model_names: list[str] = []
    machine_model_ids: list[str] = []
    ingestion_metadata_id: str | None = None

    # Cache whether the join table exists in this DB (avoids repeating to_regclass on every doc)
    global _dmm_table_exists_cache  # type: ignore
    try:
        _dmm_table_exists_cache  # type: ignore
    except Exception:
        _dmm_table_exists_cache = None  # type: ignore

    # Parse metadata_id from conventional object structure: <prefix><metadata_id>/<filename>
    rel = object_name
    if docs_prefix and rel.startswith(docs_prefix):
        rel = rel[len(docs_prefix):].lstrip("/")
    parts = [p for p in rel.split("/") if p]
    if len(parts) >= 2:
        ingestion_metadata_id = parts[0]

    if session is not None:
        try:
            from backend.utils.db import Document as DBDocument, DocumentIngestionMetadata, MachineModel

            doc = session.query(DBDocument).filter(DBDocument.gcs_path == source_gcs).first()
            if not doc and parts:
                # Best-effort fallback: match by filename (not ideal, but helps recover older rows)
                doc = session.query(DBDocument).filter(DBDocument.file_name == parts[-1]).first()

            if doc and doc.id is not None:
                # Note: DB doc.id is integer in this repo; keep as string without enforcing a specific shape.
                db_doc_id = str(doc.id)

                # Preferred: read canonical mappings from join table document_machine_models (if it exists)
                if _dmm_table_exists_cache is None:
                    try:
                        exists_row = session.execute(text("SELECT to_regclass('public.document_machine_models')")).scalar()
                        _dmm_table_exists_cache = bool(exists_row)
                    except Exception:
                        _dmm_table_exists_cache = False

                if _dmm_table_exists_cache:
                    try:
                        rows = session.execute(
                            text(
                                """
                                SELECT mm.id AS id, mm.name AS name
                                FROM public.document_machine_models dmm
                                JOIN public.machine_models mm ON mm.id = dmm.machine_model_id
                                WHERE dmm.document_id = :doc_id
                                """
                            ),
                            {"doc_id": int(doc.id)},
                        ).fetchall()
                        machine_model_ids = [str(r.id) for r in rows if getattr(r, "id", None) is not None]
                        machine_model_names = [str(r.name).strip() for r in rows if getattr(r, "name", None)]
                    except Exception as e:
                        # If the join table is missing/misconfigured, fall back to legacy field
                        logger.warning(f"Failed to query document_machine_models for document_id={doc.id}: {e}")
                        machine_model_names = _parse_machine_models(doc.machine_model)
                else:
                    # Legacy fallback
                    machine_model_names = _parse_machine_models(doc.machine_model)

            # If machine models missing, try ingestion metadata
            meta = None
            if ingestion_metadata_id:
                meta = session.query(DocumentIngestionMetadata).filter(DocumentIngestionMetadata.id == ingestion_metadata_id).first()
            if not meta:
                meta = session.query(DocumentIngestionMetadata).filter(DocumentIngestionMetadata.file_path == source_gcs).first()

            if meta and not machine_model_names:
                machine_model_names = _parse_machine_models(meta.machine_model)

            # Best-effort: resolve machine model IDs from the registry table using names
            if machine_model_names and not machine_model_ids:
                rows = session.query(MachineModel).filter(MachineModel.name.in_(machine_model_names)).all()
                machine_model_ids = [str(r.id) for r in rows if getattr(r, "id", None) is not None]
        except Exception as e:
            logger.warning(f"DB lookup failed for {source_gcs}: {e}")

    # Fallback: GCS object custom metadata
    if object_custom_metadata:
        if not machine_model_names:
            machine_model_names = _parse_machine_models(
                object_custom_metadata.get("machine_model_names")
                or object_custom_metadata.get("machineModelNames")
                or object_custom_metadata.get("machine_models")
                or object_custom_metadata.get("machineModels")
            )

        if not machine_model_ids:
            raw_ids = object_custom_metadata.get("machine_model_ids") or object_custom_metadata.get("machineModelIds")
            machine_model_ids = _parse_machine_models(raw_ids)

        if db_doc_id is None:
            raw_doc_id = object_custom_metadata.get("document_id") or object_custom_metadata.get("documentId")
            if raw_doc_id is not None:
                v = str(raw_doc_id).strip()
                db_doc_id = v if v else None

    # Deterministic fallback for document_id: UUID5(gs://bucket/object)
    if db_doc_id is None:
        fallback_id = _stable_uuid5_from_string(source_gcs)
        logger.warning(
            "Missing DB/GCS document_id; using deterministic UUID5 fallback (check DB metadata alignment!)",
            extra={"source_gcs": source_gcs, "fallback_document_id": fallback_id},
        )
        db_doc_id = str(fallback_id)

    return {
        "document_id": db_doc_id,
        "machine_model_names": machine_model_names,
        "machine_model_ids": machine_model_ids,
        "source_gcs": source_gcs,
        "ingestion_metadata_id": ingestion_metadata_id,
    }


def stage_gcs_documents_to_workdir(
    docs_bucket: str,
    docs_prefix: str,
    workdir: Path,
) -> tuple[Path, Path, list[dict[str, Any]]]:
    """
    Download all documents from GCS into workdir/documents and write workdir/doc_manifest.json.

    Returns:
        (documents_dir, manifest_path, manifest_entries)
    """
    from backend.utils.gcs_client import list_objects, download_to_path, get_object_metadata

    docs_prefix = _normalize_prefix(docs_prefix)
    documents_dir = (workdir / "documents").resolve()
    documents_dir.mkdir(parents=True, exist_ok=True)

    manifest_path = (workdir / "doc_manifest.json").resolve()

    supported = {".pdf", ".docx", ".md", ".markdown"}

    logger.info(f"Listing docs from gs://{docs_bucket}/{docs_prefix}")
    blobs = list_objects(docs_bucket, docs_prefix)

    # Best-effort DB session (optional)
    session = None
    try:
        from backend.utils.db import SessionLocal
        session = SessionLocal()
    except Exception as e:
        logger.warning(f"DB not available for metadata resolution; will use GCS metadata/fallbacks only: {e}")
        session = None

    entries: list[dict[str, Any]] = []
    try:
        for b in tqdm(blobs, desc="Staging GCS documents"):
            object_name = b.name
            if object_name.endswith("/"):
                continue
            ext = Path(object_name).suffix.lower()
            if ext not in supported:
                continue

            rel = _safe_rel_from_gcs_object(docs_prefix, object_name)
            local_path = (documents_dir / rel).resolve()
            local_path.parent.mkdir(parents=True, exist_ok=True)

            # Resolve authoritative metadata (DB → GCS custom metadata → deterministic fallback)
            custom_md = b.metadata
            if custom_md is None:
                # Only fetch metadata if we need it later; but it’s cheap enough for docs-scale runs.
                custom_md = get_object_metadata(docs_bucket, object_name)

            resolved = _resolve_authoritative_doc_metadata(
                session=session,
                docs_bucket=docs_bucket,
                docs_prefix=docs_prefix,
                object_name=object_name,
                object_custom_metadata=custom_md,
            )

            # Idempotent download: skip if local file exists and size matches
            should_download = True
            if local_path.exists() and b.size is not None:
                try:
                    if local_path.stat().st_size == int(b.size):
                        should_download = False
                except Exception:
                    should_download = True

            if should_download:
                ok = download_to_path(docs_bucket, object_name, str(local_path))
                if not ok:
                    raise RuntimeError(f"Failed to download gs://{docs_bucket}/{object_name} to {local_path}")

            entries.append(
                {
                    "document_id": resolved["document_id"],
                    # Backwards compat (existing filtering uses node.metadata["machine_model"]):
                    # - machine_models: list[str] of names
                    # - machine_model_names: list[str] of names
                    # - machine_model_ids: list[str] of ids (best-effort; may be empty)
                    "machine_models": resolved["machine_model_names"],
                    "machine_model_names": resolved["machine_model_names"],
                    "machine_model_ids": resolved["machine_model_ids"],
                    "source_gcs": resolved["source_gcs"],
                    "gcs_object_name": object_name,
                    "filename": Path(object_name).name,
                    "local_path": str(local_path),
                    "file_size_bytes": int(b.size) if b.size is not None else (local_path.stat().st_size if local_path.exists() else None),
                    "updated": b.updated.isoformat() if getattr(b, "updated", None) else None,
                    "ingestion_metadata_id": resolved.get("ingestion_metadata_id"),
                }
            )
    finally:
        try:
            if session is not None:
                session.close()
        except Exception:
            pass

    # Write manifest deterministically
    manifest_obj = {
        "docs_bucket": docs_bucket,
        "docs_prefix": docs_prefix,
        "generated_at": time.strftime("%Y-%m-%dT%H-%M-%SZ", time.gmtime()),
        "documents": entries,
    }
    with open(manifest_path, "w", encoding="utf-8") as f:
        json.dump(manifest_obj, f, indent=2, sort_keys=True)

    logger.info(f"Wrote doc manifest: {manifest_path} (documents={len(entries)})")
    return documents_dir, manifest_path, entries


def verify_local_index_artifact(index_dir: Path) -> dict[str, Any]:
    """
    Verify local index directory is present, non-empty, and contains required files + metadata.
    """
    index_dir = index_dir.resolve()
    required_files = ["docstore.json", "index_store.json", "default__vector_store.json", "index_manifest.json"]

    if not index_dir.exists() or not index_dir.is_dir():
        raise RuntimeError(f"Index directory does not exist: {index_dir}")

    missing = [f for f in required_files if not (index_dir / f).exists()]
    if missing:
        raise RuntimeError(f"Index verification failed: missing required files: {missing} in {index_dir}")

    for f in required_files:
        p = index_dir / f
        if p.stat().st_size <= 0:
            raise RuntimeError(f"Index verification failed: file is empty: {p}")

    # Validate manifest counts
    with open(index_dir / "index_manifest.json", "r", encoding="utf-8") as f:
        manifest = json.load(f) or {}
    chunk_count = manifest.get("num_chunks")
    if chunk_count is None or int(chunk_count) <= 0:
        raise RuntimeError(f"Index verification failed: num_chunks must be > 0 in index_manifest.json (got {chunk_count})")

    # Deep check: ensure required per-chunk metadata keys exist on every docstore node
    with open(index_dir / "docstore.json", "r", encoding="utf-8") as f:
        docstore = json.load(f)
    nodes = (docstore.get("docstore/data") or {})
    if not isinstance(nodes, dict) or len(nodes) == 0:
        raise RuntimeError("Index verification failed: docstore/data is empty")

    required_keys = {"document_id", "machine_models", "source_gcs", "machine_model", "machine_model_ids", "machine_model_names"}
    missing_key_counts = {k: 0 for k in required_keys}
    invalid_counts = {
        "document_id": 0,
        "machine_models": 0,
        "machine_model": 0,
        "machine_model_ids": 0,
        "machine_model_names": 0,
        "source_gcs": 0,
    }
    total = 0
    for _, wrapped in nodes.items():
        total += 1
        data = (wrapped or {}).get("__data__") if isinstance(wrapped, dict) else None
        meta = (data or {}).get("metadata") if isinstance(data, dict) else None
        if not isinstance(meta, dict):
            for k in required_keys:
                missing_key_counts[k] += 1
            continue
        for k in required_keys:
            if k not in meta:
                missing_key_counts[k] += 1

        # Validate value types/shape
        if "document_id" in meta:
            v = meta.get("document_id")
            if not (isinstance(v, str) and v.strip()):
                invalid_counts["document_id"] += 1
        if "machine_models" in meta:
            mm = meta.get("machine_models")
            if not isinstance(mm, list) or any((not isinstance(x, str)) for x in mm):
                invalid_counts["machine_models"] += 1
        if "machine_model_names" in meta:
            mmn = meta.get("machine_model_names")
            if not isinstance(mmn, list) or any((not isinstance(x, str)) for x in mmn):
                invalid_counts["machine_model_names"] += 1
        if "machine_model_ids" in meta:
            mid = meta.get("machine_model_ids")
            if not isinstance(mid, list) or any((not isinstance(x, str)) for x in mid):
                invalid_counts["machine_model_ids"] += 1
        if "machine_model" in meta:
            m = meta.get("machine_model")
            # Orchestrator supports str or list[str], but we enforce list[str] for consistency.
            if not isinstance(m, list) or any((not isinstance(x, str)) for x in m):
                invalid_counts["machine_model"] += 1
        if "source_gcs" in meta:
            sg = meta.get("source_gcs")
            if not (isinstance(sg, str) and sg.startswith("gs://")):
                invalid_counts["source_gcs"] += 1

    if any(v > 0 for v in missing_key_counts.values()) or any(v > 0 for v in invalid_counts.values()):
        raise RuntimeError(
            "Index verification failed: per-node metadata validation failed. "
            f"missing_keys={missing_key_counts} invalid_values={invalid_counts} (total_nodes={total})"
        )

    return {
        "index_dir": str(index_dir),
        "required_files": required_files,
        "num_nodes": total,
        "num_chunks": int(chunk_count),
        "manifest": manifest,
    }


def promote_index_to_gcs(
    local_index_dir: Path,
    rag_bucket: str,
    latest_prefix: str,
    old_prefix: str,
) -> dict[str, Any]:
    """
    Safe promotion: backup latest_model → verify backup → clear latest_model → upload new → verify.

    SAFETY GUARANTEE:
    - Never deletes latest_prefix until local index verification passes AND backup verification passes.
    - If upload fails, backup remains intact.
    """
    from datetime import datetime, timezone
    from backend.utils.gcs_client import list_objects, copy_prefix, delete_prefix, upload_dir, exists_prefix

    latest_prefix = _normalize_prefix(latest_prefix)
    old_prefix = _normalize_prefix(old_prefix)

    ts = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H-%M-%SZ")
    backup_prefix = f"{old_prefix}{ts}/"
    # Ensure uniqueness if rerun within the same second (idempotent backups)
    suffix = 1
    while exists_prefix(rag_bucket, backup_prefix):
        backup_prefix = f"{old_prefix}{ts}-{suffix}/"
        suffix += 1

    # Snapshot current latest contents (exact file list)
    latest_objs = list_objects(rag_bucket, latest_prefix)
    latest_names = [o.name for o in latest_objs]
    latest_rel = [n[len(latest_prefix):] if latest_prefix and n.startswith(latest_prefix) else n for n in latest_names]

    logger.info(f"[PROMOTE] Backing up gs://{rag_bucket}/{latest_prefix} -> gs://{rag_bucket}/{backup_prefix}")
    copied = copy_prefix(rag_bucket, latest_prefix, backup_prefix)

    # Verify backup: exact relative keys + count match latest snapshot
    backup_objs = list_objects(rag_bucket, backup_prefix)
    backup_names = [o.name for o in backup_objs]
    backup_rel = set([n[len(backup_prefix):] if n.startswith(backup_prefix) else n for n in backup_names])
    latest_rel_set = set(latest_rel)
    if backup_rel != latest_rel_set or len(backup_rel) != len(latest_rel):
        raise RuntimeError(
            "[PROMOTE] Backup verification failed: backup does not exactly match latest snapshot. "
            f"latest_count={len(latest_rel)} backup_count={len(backup_rel)} "
            f"missing={sorted(list(latest_rel_set - backup_rel))[:10]} "
            f"extra={sorted(list(backup_rel - latest_rel_set))[:10]}"
        )

    # Only now clear latest
    logger.info(f"[PROMOTE] Clearing gs://{rag_bucket}/{latest_prefix} (objects={len(latest_names)})")
    deleted = delete_prefix(rag_bucket, latest_prefix)

    # Upload new artifact (skip dotfiles/transient OS junk)
    logger.info(f"[PROMOTE] Uploading local index {str(local_index_dir)} -> gs://{rag_bucket}/{latest_prefix}")
    ignore_names = {".DS_Store", "Thumbs.db", ".gitkeep", ".keep"}
    uploaded = upload_dir(str(local_index_dir), rag_bucket, latest_prefix, ignore_names=ignore_names)

    # Verify remote latest matches local artifact (exact file list) and required files are present & non-empty
    local_files: set[str] = set()
    for p in Path(local_index_dir).rglob("*"):
        if p.is_file():
            if p.name in ignore_names:
                continue
            rel_parts = p.relative_to(local_index_dir).parts
            if any(part.startswith(".") for part in rel_parts):
                continue
            local_files.add(p.relative_to(local_index_dir).as_posix())

    new_latest = list_objects(rag_bucket, latest_prefix)
    remote_rel = {o.name[len(latest_prefix):] if o.name.startswith(latest_prefix) else o.name for o in new_latest}

    if remote_rel != local_files:
        raise RuntimeError(
            "[PROMOTE] Remote verification failed: remote latest_model does not match local artifact file list. "
            f"local_count={len(local_files)} remote_count={len(remote_rel)} "
            f"missing={sorted(list(local_files - remote_rel))[:10]} "
            f"extra={sorted(list(remote_rel - local_files))[:10]}"
        )

    # Runtime-required files (matches backend/rag/startup_downloader.py expectations)
    required_rel = {"index_manifest.json", "docstore.json", "index_store.json", "default__vector_store.json"}
    missing_required = sorted(list(required_rel - remote_rel))
    if missing_required:
        raise RuntimeError(f"[PROMOTE] Remote verification failed: missing required objects in latest_model: {missing_required}")

    # Ensure required objects are non-empty remotely (size>0)
    sizes_by_rel = { (o.name[len(latest_prefix):] if o.name.startswith(latest_prefix) else o.name): o.size for o in new_latest }
    empty_required = [r for r in required_rel if not sizes_by_rel.get(r) or int(sizes_by_rel.get(r) or 0) <= 0]
    if empty_required:
        raise RuntimeError(f"[PROMOTE] Remote verification failed: required objects have empty size: {empty_required}")

    return {
        "backup_prefix": backup_prefix,
        "backup_copied": copied,
        "latest_deleted": deleted,
        "uploaded": len(uploaded),
        "latest_objects": len(new_latest),
    }


def main():
    """
    Production ingestion flow:
      GCS docs -> local staging + doc_manifest.json -> chunk/embed/build local index -> verify -> (optional) promote.
    """
    # Env/config (supports both new and existing env names)
    docs_bucket = os.getenv("GCS_DOCS_BUCKET") or os.getenv("DOCS_GCS_BUCKET") or "arrow-rag-support-prod-docs"
    docs_prefix = os.getenv("GCS_DOCS_PREFIX") or os.getenv("DOCS_GCS_PREFIX") or "documents/"

    rag_bucket = os.getenv("GCS_RAG_BUCKET") or os.getenv("RAG_INDEX_GCS_BUCKET") or "arrow-rag-support-prod-rag"
    latest_prefix = os.getenv("GCS_RAG_LATEST_PREFIX") or os.getenv("RAG_INDEX_GCS_PREFIX") or "latest_model/"
    old_prefix = os.getenv("GCS_RAG_OLD_PREFIX") or "old_model/"

    default_workdir = str(Path(REPO_ROOT) / "ingest_work") if os.name == "nt" else "/workspace/ingest_work"
    workdir = Path(os.getenv("INGEST_WORKDIR", default_workdir)).resolve()

    promote = _env_bool("PROMOTE_INDEX", default=False)

    index_out_dir = (workdir / "index_artifact").resolve()
    # Keep extracted content in the same workdir for deterministic debugging
    os.environ.setdefault("EXTRACTED_CONTENT_DIR", str(workdir / "extracted_content"))

    print("\n" + "=" * 80)
    print("Arrow Production Ingestion + (Optional) Index Promotion")
    print("=" * 80)
    print(f"Docs source:  gs://{docs_bucket}/{_normalize_prefix(docs_prefix)}")
    print(f"Workdir:      {str(workdir)}")
    print(f"Index out:    {str(index_out_dir)}")
    print(f"Promote:      {promote} (PROMOTE_INDEX)")
    print(f"RAG bucket:   gs://{rag_bucket}/")
    print(f"Latest pref:  {_normalize_prefix(latest_prefix)}")
    print(f"Old pref:     {_normalize_prefix(old_prefix)}")

    workdir.mkdir(parents=True, exist_ok=True)

    # Stage docs + write manifest
    docs_dir, doc_manifest_path, entries = stage_gcs_documents_to_workdir(
        docs_bucket=docs_bucket,
        docs_prefix=docs_prefix,
        workdir=workdir,
    )

    # Build index from staged docs
    os.environ["INGEST_DOC_MANIFEST_PATH"] = str(doc_manifest_path)
    pipeline = TechnicalRAGPipeline()
    use_qdrant = _env_bool("USE_QDRANT", default=False)
    if use_qdrant:
        raise RuntimeError("Production promotion flow requires local index artifacts; USE_QDRANT must be false.")

    pipeline.build_index(
        data_dir=str(docs_dir),
        storage_dir=str(index_out_dir),
        use_qdrant=False,
    )

    # Verify local artifact
    verification = verify_local_index_artifact(index_out_dir)
    print("\n" + "=" * 80)
    print("✅ Local index verification passed")
    print(f"- Index dir: {verification['index_dir']}")
    print(f"- Num nodes: {verification['num_nodes']}")
    print(f"- Num chunks (manifest): {verification['num_chunks']}")

    # Promote to GCS (backup + swap + upload) only when PROMOTE_INDEX=true
    if promote:
        promote_result = promote_index_to_gcs(
            local_index_dir=index_out_dir,
            rag_bucket=rag_bucket,
            latest_prefix=latest_prefix,
            old_prefix=old_prefix,
        )
        print("\n" + "=" * 80)
        print("✅ Promotion completed")
        print(f"- Backup:   gs://{rag_bucket}/{promote_result['backup_prefix']}")
        print(f"- Uploaded: {promote_result['uploaded']} objects")
        print(f"- Latest objects: {promote_result['latest_objects']}")
    else:
        print("\n" + "=" * 80)
        print("ℹ️ PROMOTE_INDEX is false; skipping GCS backup/swap/upload")
        print(f"Local artifact ready at: {str(index_out_dir)}")


if __name__ == "__main__":
    main()