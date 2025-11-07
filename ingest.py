"""
High-Performance RAG Pipeline for Technical Documents
Optimized for GPU rental with bge-large-en-v1.5 and re-ranking
Enhanced with non-text content extraction (tables, images, diagrams)
"""

import warnings
# Suppress annoying Pydantic warnings
warnings.filterwarnings("ignore", message=".*validate_default.*")
warnings.filterwarnings("ignore", category=UserWarning, module="pydantic")
# Suppress pypdf warnings for malformed PDFs
warnings.filterwarnings("ignore", message=".*wrong pointing object.*")

import os
import logging
import time
import json
import yaml
import re
from typing import List, Dict, Any, Optional, Tuple
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path
import base64
from io import BytesIO

import fitz  # PyMuPDF
import pandas as pd
from PIL import Image
import numpy as np
from tqdm import tqdm

from llama_index.core import SimpleDirectoryReader, VectorStoreIndex, StorageContext, load_index_from_storage, Settings
from llama_index.core.node_parser import SentenceSplitter
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from llama_index.core.schema import NodeWithScore, TextNode, ImageNode, Document

# DOCX and Markdown support
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
from llama_index.vector_stores.qdrant import QdrantVectorStore
from llama_index.core.storage.docstore import SimpleDocumentStore
from llama_index.core.storage.index_store import SimpleIndexStore
from sentence_transformers import CrossEncoder
import qdrant_client
import shutil
import tarfile

# Try to import Anthropic (optional dependency)
try:
    from anthropic import Anthropic
    ANTHROPIC_AVAILABLE = True
except ImportError:
    ANTHROPIC_AVAILABLE = False
    logger.warning("Anthropic package not available. Claude rewriting will be disabled.")

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Suppress pypdf warnings for malformed PDFs
logging.getLogger("pypdf").setLevel(logging.ERROR)
logging.getLogger("pypdf._reader").setLevel(logging.ERROR)

# Log DOCX availability after logger is initialized
if not DOCX_AVAILABLE:
    logger.warning("python-docx not available. DOCX support will be disabled.")


class TextPreprocessor:
    """
    Enhanced AI-powered text preprocessor for RAG pipeline.
    Removes boilerplate, normalizes technical content, fixes artifacts, and filters low-quality chunks.
    """
    
    def __init__(self):
        # Compile regex patterns for common boilerplate text
        # Using more specific patterns to avoid catastrophic backtracking
        self.boilerplate_patterns = [
            # Page numbers (various formats) - more specific to avoid hanging
            re.compile(r'^[ \t]*[Pp]age[ \t]+\d+[ \t]*$', re.MULTILINE),
            re.compile(r'^[ \t]*\d{1,4}[ \t]*$', re.MULTILINE),  # Standalone page numbers (limit digits)
            re.compile(r'^[ \t]*-[ \t]*\d+[ \t]*-[ \t]*$', re.MULTILINE),  # "- 5 -"
            
            # Confidential/Proprietary markings
            re.compile(r'\b[Mm]emjet[ \t]+[Cc]onfidential\b', re.IGNORECASE),
            re.compile(r'\b[Cc]onfidential\b', re.IGNORECASE),
            re.compile(r'\b[Pp]roprietary\b', re.IGNORECASE),
            
            # Copyright notices (various formats) - limit length to prevent hanging
            re.compile(r'©[ \t]*\d{4}[^\n]{0,100}$', re.MULTILINE | re.IGNORECASE),
            re.compile(r'Copyright[ \t]+©?[ \t]*\d{4}[^\n]{0,100}$', re.MULTILINE | re.IGNORECASE),
            re.compile(r'All[ \t]+rights[ \t]+reserved\.?', re.IGNORECASE),
            
            # Common header/footer patterns
            re.compile(r'^[ \t]*(Document|Version|Rev|Revision)[ \t]*:[ \t]*[\w\.-]+[ \t]*$', re.MULTILINE | re.IGNORECASE),
            
            # Repeated section titles (if they appear multiple times)
            re.compile(r'^[ \t]*(Table[ \t]+of[ \t]+Contents|Contents|Index)[ \t]*$', re.MULTILINE | re.IGNORECASE),
            
            # Date stamps in headers/footers
            re.compile(r'^[ \t]*\d{1,2}[/-]\d{1,2}[/-]\d{2,4}[ \t]*$', re.MULTILINE),
        ]
        
        # Enhanced header/footer patterns (more comprehensive)
        self.header_footer_patterns = [
            re.compile(r'^[ \t]*(DuraFlex|DuraCore|DuraBolt|anyCUT|EZCut)[ \t]+.*?[ \t]*$', re.MULTILINE | re.IGNORECASE),
            re.compile(r'^[ \t]*[A-Z][a-z]+[ \t]+(Manual|Guide|Databook|Release Notes)[ \t]*$', re.MULTILINE),
            re.compile(r'^[ \t]*V\d+\.\d+[ \t]*$', re.MULTILINE),  # Version numbers
            re.compile(r'^[ \t]*Rev[ \t]*\d+[ \t]*$', re.MULTILINE | re.IGNORECASE),
        ]
        
        # Table of Contents detection patterns
        self.toc_patterns = [
            re.compile(r'^[ \t]*(Table[ \t]+of[ \t]+Contents|Contents|Index|TOC)[ \t]*$', re.MULTILINE | re.IGNORECASE),
            re.compile(r'^\s*\d+\.\d+[ \t]+.*?\s+\d+$', re.MULTILINE),  # "1.2 Section Name    5"
            re.compile(r'^\s*[A-Z][a-z]+[ \t]+\.{3,}[ \t]+\d+$', re.MULTILINE),  # "Section .......... 10"
        ]
        
        # Patterns for structured lines that should be preserved (for smart chunking)
        self.preserve_patterns = [
            re.compile(r'^(Usage|Command|Example|Syntax|Parameters?|Options?|Steps?|Procedure|Note|Warning|Important):\s*', re.MULTILINE | re.IGNORECASE),
            re.compile(r'^\s*\d+[\.\)]\s+', re.MULTILINE),  # Numbered lists: "1. ", "2) "
            re.compile(r'^[-*•]\s+', re.MULTILINE),  # Bullet points
            re.compile(r'^\s*[A-Z][a-z]+:\s*$', re.MULTILINE),  # Section headers ending with colon
        ]
        
        # Common redundant phrases in technical docs
        self.redundant_phrases = [
            (re.compile(r'\bplease[ \t]+note[ \t]+that\b', re.IGNORECASE), ''),
            (re.compile(r'\bit[ \t]+is[ \t]+important[ \t]+to[ \t]+note[ \t]+that\b', re.IGNORECASE), ''),
            (re.compile(r'\bas[ \t]+you[ \t]+can[ \t]+see\b', re.IGNORECASE), ''),
            (re.compile(r'\bas[ \t]+shown[ \t]+above\b', re.IGNORECASE), ''),
            (re.compile(r'\bas[ \t]+shown[ \t]+below\b', re.IGNORECASE), ''),
            (re.compile(r'\bas[ \t]+mentioned[ \t]+previously\b', re.IGNORECASE), ''),
            (re.compile(r'\bas[ \t]+mentioned[ \t]+earlier\b', re.IGNORECASE), ''),
            (re.compile(r'\bfor[ \t]+more[ \t]+information[ \t]+please[ \t]+refer[ \t]+to\b', re.IGNORECASE), 'See'),
            (re.compile(r'\bfor[ \t]+additional[ \t]+details[ \t]+please[ \t]+see\b', re.IGNORECASE), 'See'),
        ]
    
    def is_table_of_contents(self, text: str) -> bool:
        """
        Detect if text is a Table of Contents section.
        Returns True if TOC patterns are found.
        """
        if not text:
            return False
        
        lines = text.split('\n')
        toc_line_count = 0
        
        # Check for TOC header
        for pattern in self.toc_patterns[:1]:  # First pattern is TOC header
            if pattern.search(text):
                toc_line_count += 1
        
        # Check for TOC entry patterns (section numbers with page numbers)
        for line in lines[:20]:  # Check first 20 lines
            for pattern in self.toc_patterns[1:]:
                if pattern.search(line):
                    toc_line_count += 1
                    break
        
        # If we find TOC header + multiple TOC entries, it's likely a TOC
        return toc_line_count >= 3
    
    def remove_table_of_contents(self, text: str) -> str:
        """Remove Table of Contents sections from text."""
        if not self.is_table_of_contents(text):
            return text
        
        lines = text.split('\n')
        cleaned_lines = []
        in_toc = False
        
        for line in lines:
            # Check if line starts TOC
            if any(pattern.search(line) for pattern in self.toc_patterns[:1]):
                in_toc = True
                continue
            
            # Check if line is TOC entry
            if in_toc:
                if any(pattern.search(line) for pattern in self.toc_patterns[1:]):
                    continue
                # Check if we've left TOC (found non-TOC content)
                if line.strip() and not any(pattern.search(line) for pattern in self.toc_patterns):
                    # Look ahead: if next few lines aren't TOC entries, we've left TOC
                    in_toc = False
            
            if not in_toc:
                cleaned_lines.append(line)
        
        return '\n'.join(cleaned_lines)
    
    def remove_headers_footers(self, text: str) -> str:
        """Remove header and footer text that appears on multiple pages."""
        if not text:
            return text
        
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            # Skip header/footer patterns
            is_header_footer = False
            
            # Check enhanced header/footer patterns
            for pattern in self.header_footer_patterns:
                if pattern.search(line):
                    is_header_footer = True
                    break
            
            # Check if line is very short and appears to be header/footer
            if not is_header_footer and len(line.strip()) < 50:
                # Check for common header/footer indicators
                if (line.strip().count(' ') < 5 and 
                    (line.strip().isupper() or 
                     re.match(r'^[A-Z][a-z]+[ \t]+(Manual|Guide|V\d+|Rev)', line.strip()))):
                    is_header_footer = True
            
            if not is_header_footer:
                cleaned_lines.append(line)
        
        return '\n'.join(cleaned_lines)
    
    def is_first_page_without_content(self, text: str, metadata: dict = None) -> bool:
        """
        Detect if this is a first page (cover page) with no meaningful content.
        Checks for title pages, cover pages, etc.
        """
        if not text:
            return True
        
        # Check page number
        page_label = metadata.get('page_label', '') if metadata else ''
        if page_label and page_label not in ['1', 'i', 'I']:
            return False
        
        # Check word count
        words = len(text.split())
        if words < 15:
            return True
        
        # Check for cover page indicators
        cover_indicators = [
            r'^[A-Z][A-Z\s]{10,}$',  # All caps title
            r'^(User|Installation|Service|Operation)[ \t]+(Manual|Guide|Databook)',  # Title format
        ]
        
        lines = text.split('\n')[:10]  # Check first 10 lines
        cover_line_count = 0
        
        for line in lines:
            for pattern in cover_indicators:
                if re.match(pattern, line.strip(), re.IGNORECASE):
                    cover_line_count += 1
        
        # If mostly cover page content and low word count
        return cover_line_count >= 2 and words < 50
    
    def fix_hyphenation(self, text: str) -> str:
        """
        Fix hyphenated words split across lines.
        Example: "print-\nhead" -> "printhead"
        """
        if not text:
            return text
        
        # Pattern: word ending with hyphen, followed by newline, followed by word continuation
        # Match: "word-\nword" -> "wordword"
        text = re.sub(r'([a-zA-Z])-\s*\n\s*([a-zA-Z])', r'\1\2', text)
        
        # Also handle cases with spaces: "word- \n word"
        text = re.sub(r'([a-zA-Z])-\s+\n\s+([a-zA-Z])', r'\1\2', text)
        
        return text
    
    def fix_line_breaks(self, text: str) -> str:
        """
        Fix inappropriate line breaks in the middle of sentences.
        Preserves intentional paragraph breaks (double newlines).
        """
        if not text:
            return text
        
        lines = text.split('\n')
        fixed_lines = []
        i = 0
        
        while i < len(lines):
            line = lines[i].strip()
            
            if not line:
                fixed_lines.append('')
                i += 1
                continue
            
            # Check if line ends with sentence-ending punctuation
            ends_with_punctuation = re.search(r'[.!?]\s*$', line)
            
            # Check if next line starts with capital letter (new sentence)
            next_starts_sentence = False
            if i + 1 < len(lines):
                next_line = lines[i + 1].strip()
                next_starts_sentence = bool(re.match(r'^[A-Z]', next_line))
            
            # If line doesn't end with punctuation and next doesn't start sentence,
            # it's likely a broken line - join them
            if not ends_with_punctuation and not next_starts_sentence and i + 1 < len(lines):
                next_line = lines[i + 1].strip()
                if next_line and not next_line.startswith(('•', '-', '*', '1.', '2.', '3.')):
                    # Join with space
                    line = line + ' ' + next_line
                    i += 1  # Skip next line since we joined it
            
            fixed_lines.append(line)
            i += 1
        
        return '\n'.join(fixed_lines)
    
    def remove_repeated_phrases(self, text: str) -> str:
        """Remove redundant phrases that don't add semantic value."""
        cleaned = text
        
        for pattern, replacement in self.redundant_phrases:
            try:
                cleaned = pattern.sub(replacement, cleaned)
            except Exception as e:
                logger.debug(f"Failed to remove redundant phrase: {e}")
                continue
        
        return cleaned
    
    def normalize_technical_content(self, text: str) -> str:
        """
        Normalize technical instructions and explanations.
        Fixes spacing, punctuation, and formatting issues.
        """
        if not text:
            return text
        
        # Fix spacing around punctuation
        text = re.sub(r'\s+([.,;:!?])', r'\1', text)  # Remove space before punctuation
        text = re.sub(r'([.,;:!?])([^\s])', r'\1 \2', text)  # Add space after punctuation
        
        # Fix spacing around parentheses and brackets
        text = re.sub(r'\(\s+', '(', text)
        text = re.sub(r'\s+\)', ')', text)
        text = re.sub(r'\[\s+', '[', text)
        text = re.sub(r'\s+\]', ']', text)
        
        # Normalize multiple spaces (but preserve intentional spacing in tables/code)
        # Only normalize if not in a table-like structure
        if '|' not in text and '\t' not in text:
            text = re.sub(r' {2,}', ' ', text)
        
        # Fix common technical formatting issues
        text = re.sub(r'(\d+)\s*-\s*(\d+)', r'\1-\2', text)  # Number ranges: "5 - 10" -> "5-10"
        text = re.sub(r'(\w+)\s*/\s*(\w+)', r'\1/\2', text)  # Slashes: "A / B" -> "A/B"
        
        return text
    
    def remove_boilerplate(self, text: str) -> str:
        """Remove common boilerplate text patterns from the input."""
        cleaned = text
        
        # Apply each boilerplate pattern with error handling
        for pattern in self.boilerplate_patterns:
            try:
                cleaned = pattern.sub('', cleaned)
            except Exception as e:
                # If regex fails, skip this pattern and continue
                logger.debug(f"Regex pattern failed, skipping: {e}")
                continue
        
        return cleaned
    
    def normalize_whitespace(self, text: str) -> str:
        """Normalize whitespace: collapse multiple spaces/tabs, normalize newlines."""
        try:
            # Replace tabs with spaces
            text = text.replace('\t', ' ')
            
            # Collapse multiple spaces into single space (limit to prevent hanging)
            text = re.sub(r' {2,}', ' ', text)
            
            # Normalize line breaks: multiple newlines -> double newline (paragraph break)
            text = re.sub(r'\n{3,}', '\n\n', text)
            
            # Remove leading/trailing whitespace from each line
            lines = [line.strip() for line in text.split('\n')]
            text = '\n'.join(lines)
            
            # Remove leading/trailing whitespace from entire text
            text = text.strip()
        except Exception as e:
            logger.debug(f"Whitespace normalization failed: {e}")
            # Return original text if normalization fails
            pass
        
        return text
    
    def should_preserve_line(self, line: str) -> bool:
        """Check if a line matches patterns that should be preserved as-is."""
        for pattern in self.preserve_patterns:
            if pattern.search(line):
                return True
        return False
    
    def clean_text(self, text: str, metadata: dict = None) -> str:
        """
        Apply all enhanced cleaning steps in optimal order.
        Returns cleaned text ready for chunking and embedding.
        """
        if not text:
            return text
        
        cleaned = text
        
        # Step 1: Remove Table of Contents
        cleaned = self.remove_table_of_contents(cleaned)
        
        # Step 2: Remove headers and footers
        cleaned = self.remove_headers_footers(cleaned)
        
        # Step 3: Remove boilerplate (copyright, page numbers, etc.)
        cleaned = self.remove_boilerplate(cleaned)
        
        # Step 4: Fix text artifacts
        cleaned = self.fix_hyphenation(cleaned)
        cleaned = self.fix_line_breaks(cleaned)
        
        # Step 5: Remove redundant phrases
        cleaned = self.remove_repeated_phrases(cleaned)
        
        # Step 6: Normalize technical content (spacing, punctuation)
        cleaned = self.normalize_technical_content(cleaned)
        
        # Step 7: Normalize whitespace (final pass)
        cleaned = self.normalize_whitespace(cleaned)
        
        return cleaned
    
    def is_low_content_page(self, text: str, min_words: int = 15) -> bool:
        """Check if a page has too little content to be useful."""
        words = len(text.split())
        return words < min_words
    
    def should_skip_node(self, text: str, min_chars: int = 30, metadata: dict = None) -> Tuple[bool, str]:
        """
        Check if a node should be skipped (too short or empty).
        Returns (should_skip: bool, reason: str)
        """
        if not text:
            return True, "empty_text"
        
        text_stripped = text.strip()
        if len(text_stripped) < min_chars:
            return True, "too_short"
        
        # Check if it's mostly whitespace or special characters
        alpha_chars = len(re.findall(r'[a-zA-Z]', text))
        if alpha_chars < min_chars // 2:  # At least half should be alphabetic
            return True, "low_alphabetic_content"
        
        # Check if it's a Table of Contents
        if self.is_table_of_contents(text):
            return True, "table_of_contents"
        
        # Check if it's a first page without content
        if self.is_first_page_without_content(text, metadata):
            return True, "first_page_no_content"
        
        return False, ""


class ClaudeSemanticRewriter:
    """
    Uses Claude API to semantically rewrite text chunks for improved clarity
    while preserving technical meaning and structured content.
    """
    
    def __init__(self, api_key: Optional[str] = None, model: str = "claude-3-5-sonnet-20241022", 
                 enabled: bool = False, max_retries: int = 2, timeout: int = 30):
        """
        Initialize Claude semantic rewriter.
        
        Args:
            api_key: Anthropic API key (defaults to ANTHROPIC_API_KEY env var)
            model: Claude model to use (default: claude-3-5-sonnet-20241022)
            enabled: Whether rewriting is enabled (default: False)
            max_retries: Maximum retry attempts per chunk
            timeout: Request timeout in seconds
        """
        self.enabled = enabled and ANTHROPIC_AVAILABLE
        self.model = model
        self.max_retries = max_retries
        self.timeout = timeout
        self.client = None
        
        if self.enabled:
            api_key = api_key or os.getenv('ANTHROPIC_API_KEY')
            if not api_key:
                logger.warning("⚠️ Claude rewriting enabled but ANTHROPIC_API_KEY not found. Disabling rewriting.")
                self.enabled = False
            else:
                try:
                    self.client = Anthropic(api_key=api_key, timeout=timeout)
                    logger.info(f"✅ Claude semantic rewriter initialized (model: {model})")
                except Exception as e:
                    logger.warning(f"⚠️ Failed to initialize Claude client: {e}. Disabling rewriting.")
                    self.enabled = False
    
    def _is_structured_content(self, text: str, metadata: dict) -> bool:
        """
        Check if content is structured (table, code, list) that should be preserved as-is.
        """
        content_type = metadata.get("content_type", "text")
        
        # Don't rewrite tables, images, or captions
        if content_type in ["table", "image", "figure_caption"]:
            return True
        
        # Check for code blocks
        if '```' in text or '`' in text:
            return True
        
        # Check for markdown tables
        if text.count('|') >= 6 and re.search(r'\|[^\|]+\|', text):
            return True
        
        # Check for dense lists (many bullets/numbers)
        lines = text.split('\n')
        list_lines = sum(1 for line in lines if re.match(r'^\s*[-*•\d]', line))
        if list_lines >= len(lines) * 0.5:  # 50% or more are list items
            return True
        
        return False
    
    def _create_rewrite_prompt(self, text: str, metadata: dict) -> str:
        """Create prompt for Claude to rewrite the text."""
        content_type = metadata.get("content_type", "text")
        file_name = metadata.get("file_name", "document")
        
        prompt = f"""Rewrite the following technical documentation text to improve semantic clarity while preserving all technical meaning and accuracy.

**Requirements:**
1. Keep all technical terms, specifications, and measurements exactly as written
2. Remove minor redundancies, filler phrases, and ambiguous wording
3. Improve sentence flow and clarity
4. Preserve all structured content (tables, lists, code blocks) exactly as-is
5. Maintain the same level of technical detail
6. Do not add new information or remove important details
7. Output ONLY the rewritten text, no explanations or markdown formatting

**Source:** {file_name}
**Content Type:** {content_type}

**Original Text:**
{text}

**Rewritten Text:**"""
        
        return prompt
    
    def rewrite_chunk(self, node: TextNode) -> Tuple[TextNode, bool]:
        """
        Rewrite a single chunk using Claude API.
        
        Args:
            node: TextNode to rewrite
            
        Returns:
            Tuple of (rewritten_node, was_rewritten: bool)
        """
        if not self.enabled or not self.client:
            return node, False
        
        # Skip structured content
        if self._is_structured_content(node.text, node.metadata):
            logger.debug(f"Skipping structured content rewrite: {node.metadata.get('file_name', 'unknown')}")
            return node, False
        
        # Skip very short chunks (not worth API call)
        if len(node.text.strip()) < 100:
            return node, False
        
        # Create prompt
        prompt = self._create_rewrite_prompt(node.text, node.metadata)
        
        # Call Claude API with retries
        for attempt in range(self.max_retries):
            try:
                response = self.client.messages.create(
                    model=self.model,
                    max_tokens=4096,
                    temperature=0.1,  # Low temperature for consistency
                    messages=[{
                        "role": "user",
                        "content": prompt
                    }]
                )
                
                # Extract rewritten text
                rewritten_text = response.content[0].text.strip()
                
                # Validate: rewritten text should not be empty and should be reasonable length
                if not rewritten_text or len(rewritten_text) < len(node.text) * 0.5:
                    logger.warning(f"Claude rewrite too short or empty, using original. Chunk: {node.metadata.get('file_name', 'unknown')}")
                    return node, False
                
                # Create new node with rewritten text but same metadata
                rewritten_node = TextNode(
                    text=rewritten_text,
                    metadata=node.metadata.copy()  # Preserve all metadata
                )
                
                logger.debug(f"Successfully rewritten chunk from {node.metadata.get('file_name', 'unknown')}")
                return rewritten_node, True
                
            except Exception as e:
                if attempt < self.max_retries - 1:
                    logger.warning(f"Claude rewrite attempt {attempt + 1} failed, retrying: {e}")
                    time.sleep(1)  # Brief delay before retry
                else:
                    logger.warning(f"Claude rewrite failed after {self.max_retries} attempts, using original: {e}")
                    return node, False
        
        return node, False
    
    def rewrite_nodes(self, nodes: List[TextNode], show_progress: bool = True) -> Tuple[List[TextNode], Dict[str, int]]:
        """
        Rewrite a list of TextNodes using Claude API.
        
        Args:
            nodes: List of TextNode objects to rewrite
            show_progress: Whether to show progress bar
            
        Returns:
            Tuple of (rewritten_nodes, stats_dict)
        """
        if not self.enabled:
            logger.info("Claude rewriting is disabled, skipping rewrite step")
            return nodes, {"rewritten": 0, "skipped": len(nodes), "failed": 0, "structured": 0}
        
        logger.info(f"🔄 Starting Claude semantic rewriting for {len(nodes)} chunks...")
        
        rewritten_nodes = []
        stats = {
            "rewritten": 0,
            "skipped": 0,
            "failed": 0,
            "structured": 0
        }
        
        # Process nodes with progress bar
        iterator = tqdm(nodes, desc="   Rewriting chunks", disable=not show_progress) if show_progress else nodes
        
        for node in iterator:
            # Check if structured content
            if self._is_structured_content(node.text, node.metadata):
                rewritten_nodes.append(node)
                stats["structured"] += 1
                continue
            
            # Rewrite the chunk
            rewritten_node, was_rewritten = self.rewrite_chunk(node)
            rewritten_nodes.append(rewritten_node)
            
            if was_rewritten:
                stats["rewritten"] += 1
            else:
                # Determine why it wasn't rewritten
                if not self.enabled:
                    stats["skipped"] += 1
                elif len(node.text.strip()) < 100:
                    stats["skipped"] += 1
                else:
                    stats["failed"] += 1
        
        logger.info(f"✅ Claude rewriting complete: {stats['rewritten']} rewritten, {stats['structured']} structured (preserved), "
                   f"{stats['skipped']} skipped, {stats['failed']} failed")
        
        return rewritten_nodes, stats


class SmartChunkSplitter:
    """
    Wrapper around SentenceSplitter that preserves structured content like tables,
    code blocks, numbered steps, and command syntax.
    """
    
    def __init__(self, chunk_size: int = 350, chunk_overlap: int = 88, preprocessor: TextPreprocessor = None):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.base_splitter = SentenceSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            include_metadata=True
        )
        self.preprocessor = preprocessor or TextPreprocessor()
    
    def _is_table_content(self, text: str) -> bool:
        """Detect if text contains table-like content (markdown table or pipe-separated)."""
        # Check for markdown table pattern (| col1 | col2 |)
        if re.search(r'\|[^\|]+\|', text) and text.count('|') >= 6:
            return True
        # Check for tab-separated or multiple spaces (likely table)
        lines = text.split('\n')
        if len(lines) >= 3:
            tabs_or_spaces = sum(1 for line in lines if '\t' in line or re.search(r' {3,}', line))
            if tabs_or_spaces >= len(lines) * 0.6:  # 60% of lines have table-like spacing
                return True
        return False
    
    def _is_code_block(self, text: str) -> bool:
        """Detect if text contains code blocks."""
        # Check for code block markers
        if '```' in text or '`' in text:
            return True
        # Check for common code patterns
        code_keywords = ['def ', 'class ', 'import ', 'function', 'return', 'const ', 'var ', 'let ']
        if any(keyword in text for keyword in code_keywords):
            return True
        return False
    
    def _preserve_structured_chunks(self, text: str) -> List[str]:
        """
        Split text while preserving structured content as single units.
        Returns list of text chunks.
        """
        chunks = []
        lines = text.split('\n')
        current_chunk = []
        current_chunk_size = 0
        
        i = 0
        while i < len(lines):
            line = lines[i]
            line_stripped = line.strip()
            
            # Check if this line starts a structured block
            if self.preprocessor.should_preserve_line(line_stripped):
                # Try to collect the entire structured block
                structured_block = [line]
                block_size = len(line)
                j = i + 1
                
                # Collect lines until we hit a non-structured line or exceed chunk size
                while j < len(lines) and block_size + len(lines[j]) < self.chunk_size:
                    next_line = lines[j].strip()
                    # Continue if it's part of the structure (numbered, bullet, or continuation)
                    if (next_line.startswith((' ', '\t')) or  # Indented continuation
                        re.match(r'^\d+[\.\)]', next_line) or  # Next numbered item
                        re.match(r'^[-*•]', next_line) or  # Next bullet
                        re.match(r'^[A-Z][a-z]+:\s*$', next_line)):  # Next section header
                        structured_block.append(lines[j])
                        block_size += len(lines[j])
                        j += 1
                    else:
                        break
                
                # If we have accumulated text, save it first
                if current_chunk:
                    chunks.append('\n'.join(current_chunk))
                    current_chunk = []
                    current_chunk_size = 0
                
                # Add the structured block as a single chunk
                structured_text = '\n'.join(structured_block)
                # Check if it exceeds chunk size (rare, but handle it)
                if len(structured_text) > self.chunk_size:
                    # Split the structured block using base splitter, but preserve internal structure
                    sub_chunks = self.base_splitter.split_text(structured_text)
                    chunks.extend(sub_chunks)
                else:
                    chunks.append(structured_text)
                
                i = j
                continue
            
            # Regular line - add to current chunk
            line_size = len(line)
            if current_chunk_size + line_size <= self.chunk_size:
                current_chunk.append(line)
                current_chunk_size += line_size
                i += 1
            else:
                # Current chunk is full, save it
                if current_chunk:
                    chunks.append('\n'.join(current_chunk))
                    # Start new chunk with overlap
                    overlap_lines = []
                    overlap_size = 0
                    # Get last few lines for overlap (up to chunk_overlap chars)
                    for line_idx in range(len(current_chunk) - 1, -1, -1):
                        if overlap_size + len(current_chunk[line_idx]) <= self.chunk_overlap:
                            overlap_lines.insert(0, current_chunk[line_idx])
                            overlap_size += len(current_chunk[line_idx])
                        else:
                            break
                    current_chunk = overlap_lines + [line]
                    current_chunk_size = overlap_size + line_size
                    i += 1
        
        # Add remaining chunk
        if current_chunk:
            chunks.append('\n'.join(current_chunk))
        
        return chunks
    
    def split_text(self, text: str) -> List[str]:
        """Split text with smart chunking that preserves structured content."""
        # Check if entire text is a table or code block
        if self._is_table_content(text) or self._is_code_block(text):
            # Preserve as single chunk (may exceed chunk_size, but that's okay for structured content)
            return [text]
        
        # Use smart chunking
        return self._preserve_structured_chunks(text)
    
    def get_nodes_from_documents(self, documents: List[Document], show_progress: bool = False) -> List[TextNode]:
        """
        Split documents into nodes with smart chunking.
        This is a simplified version that processes documents one by one.
        """
        all_nodes = []
        
        for doc in tqdm(documents, desc="Splitting documents", disable=not show_progress):
            try:
                text = doc.text or ""
                doc_name = doc.metadata.get('file_name', 'unknown')
                
                # Clean the text first (with error handling)
                try:
                    text = self.preprocessor.clean_text(text, metadata=doc.metadata)
                except Exception as e:
                    logger.warning(f"Error cleaning text for {doc_name}: {e}")
                    # Use original text if cleaning fails
                    text = doc.text or ""
                
                # Check if page/document should be skipped (low content)
                try:
                    if self.preprocessor.is_low_content_page(text):
                        logger.debug(f"Skipping low-content page: {doc_name}")
                        continue
                except Exception as e:
                    logger.debug(f"Error checking low-content for {doc_name}: {e}")
                    # Continue processing if check fails
                
                # Split into chunks (with error handling)
                try:
                    chunks = self.split_text(text)
                except Exception as e:
                    logger.error(f"Error splitting text for {doc_name}: {e}")
                    # Fallback to simple split if smart chunking fails
                    if text:
                        chunks = [text]
                    else:
                        chunks = []
                
                # Create nodes from chunks
                for chunk_idx, chunk_text in enumerate(chunks):
                    # Enhanced skip check with reason tracking
                    should_skip, skip_reason = self.preprocessor.should_skip_node(
                        chunk_text, 
                        metadata={**doc.metadata, "chunk_index": chunk_idx, "total_chunks": len(chunks)}
                    )
                    
                    if should_skip:
                        logger.debug(f"Skipping chunk {chunk_idx} from {doc_name}: {skip_reason}")
                        continue
                    
                    # Create node with metadata (including skip reason if applicable)
                    try:
                        node = TextNode(
                            text=chunk_text,
                            metadata={
                                **doc.metadata,
                                "chunk_index": chunk_idx,
                                "total_chunks": len(chunks),
                                "content_type": "text"
                            }
                        )
                        all_nodes.append(node)
                    except Exception as e:
                        logger.error(f"Error creating node for {doc_name}, chunk {chunk_idx}: {e}")
                        continue
                        
            except Exception as e:
                logger.error(f"Error processing document {doc.metadata.get('file_name', 'unknown')}: {e}")
                continue
        
        return all_nodes


class NonTextExtractor:
    """Extract and process non-text content from documents."""
    
    def __init__(self, output_dir="/workspace/extracted_content"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        
    def extract_tables_from_pdf(self, pdf_path: str) -> List[Dict[str, Any]]:
        """Extract tables from PDF using PyMuPDF."""
        tables = []
        doc = fitz.open(pdf_path)
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            # Extract tables using PyMuPDF's table detection
            page_tables = page.find_tables()
            
            for table_idx, table in enumerate(page_tables):
                try:
                    # Extract table data
                    table_data = table.extract()
                    if table_data and len(table_data) > 1:  # Ensure we have headers and data
                        # Convert to pandas DataFrame for better structure
                        headers = table_data[0]
                        # Handle duplicate column names
                        unique_headers = []
                        for i, header in enumerate(headers):
                            if header in unique_headers:
                                unique_headers.append(f"{header}_{i}")
                            else:
                                unique_headers.append(header)
                        
                        df = pd.DataFrame(table_data[1:], columns=unique_headers)
                        
                        # Create table metadata
                        table_info = {
                            "source_path": pdf_path,
                            "page_number": page_num + 1,
                            "table_index": table_idx,
                            "table_data": df.to_dict('records'),
                            "table_markdown": df.to_markdown(index=False),
                            "table_json": df.to_json(orient='records'),
                            "row_count": len(df),
                            "column_count": len(df.columns),
                            "content_type": "table"
                        }
                        tables.append(table_info)
                        
                        # Save table as separate file
                        table_filename = f"{Path(pdf_path).stem}_page{page_num+1}_table{table_idx}.json"
                        table_path = self.output_dir / table_filename
                        with open(table_path, 'w', encoding='utf-8') as f:
                            json.dump(table_info, f, indent=2, ensure_ascii=False)
                            
                except Exception as e:
                    logger.warning(f"Failed to extract table {table_idx} from page {page_num + 1}: {e}")
                    continue
                    
        doc.close()
        return tables
    
    def extract_images_from_pdf(self, pdf_path: str) -> List[Dict[str, Any]]:
        """Extract images and diagrams from PDF."""
        images = []
        doc = fitz.open(pdf_path)
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            image_list = page.get_images()
            
            for img_idx, img in enumerate(image_list):
                try:
                    # Get image data
                    xref = img[0]
                    pix = fitz.Pixmap(doc, xref)
                    
                    if pix.n - pix.alpha < 4:  # GRAY or RGB
                        # Convert to PIL Image
                        img_data = pix.tobytes("png")
                        pil_image = Image.open(BytesIO(img_data))
                        
                        # Get image metadata
                        img_rects = page.get_image_rects(xref)
                        img_rect = img_rects[0] if img_rects else None
                        
                        # Create image info (NO base64 embedding - only metadata)
                        image_info = {
                            "source_path": pdf_path,
                            "page_number": page_num + 1,
                            "image_index": img_idx,
                            # Removed: image_data base64 encoding (not needed for embeddings)
                            "width": pil_image.width,
                            "height": pil_image.height,
                            "format": "PNG",
                            "content_type": "image",
                            "caption": f"Image from {Path(pdf_path).stem}, page {page_num + 1}",
                            "bbox": str(img_rect) if img_rect else None
                        }
                        images.append(image_info)
                        
                        # Save image
                        img_filename = f"{Path(pdf_path).stem}_page{page_num+1}_img{img_idx}.png"
                        img_path = self.output_dir / img_filename
                        pil_image.save(img_path)
                        image_info["saved_path"] = str(img_path)
                        
                except Exception as e:
                    logger.warning(f"Failed to extract image {img_idx} from page {page_num + 1}: {e}")
                    continue
                    
        doc.close()
        return images
    
    def extract_figure_captions(self, pdf_path: str) -> List[Dict[str, Any]]:
        """Extract figure captions and references."""
        captions = []
        doc = fitz.open(pdf_path)
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text()
            
            # Look for figure captions (simple pattern matching)
            lines = text.split('\n')
            for i, line in enumerate(lines):
                line_lower = line.lower().strip()
                if any(keyword in line_lower for keyword in ['figure', 'fig.', 'diagram', 'chart', 'graph']):
                    # Extract caption text
                    caption_text = line.strip()
                    if len(caption_text) > 10:  # Filter out very short matches
                        caption_info = {
                            "source_path": pdf_path,
                            "page_number": page_num + 1,
                            "caption_text": caption_text,
                            "content_type": "figure_caption",
                            "line_number": i + 1
                        }
                        captions.append(caption_info)
        
        doc.close()
        return captions


class DocumentLoader:
    """
    Custom document loader that supports PDF, DOCX, and Markdown files.
    Preserves document provenance with file_name and page_label/section metadata.
    """
    
    def __init__(self, data_dir: str):
        self.data_dir = Path(data_dir)
        self.supported_extensions = {'.pdf', '.docx', '.md', '.markdown'}
    
    def load_documents(self) -> List[Document]:
        """
        Load all supported documents from data directory.
        Returns list of Document objects with proper metadata.
        """
        documents = []
        
        # Get all supported files
        all_files = []
        for ext in self.supported_extensions:
            all_files.extend(list(self.data_dir.glob(f"**/*{ext}")))
        
        logger.info(f"Found {len(all_files)} files to process")
        
        for file_path in tqdm(all_files, desc="Loading documents"):
            try:
                file_ext = file_path.suffix.lower()
                file_name = file_path.name
                
                if file_ext == '.pdf':
                    # Use SimpleDirectoryReader for PDFs (existing logic)
                    pdf_docs = SimpleDirectoryReader(input_files=[str(file_path)]).load_data()
                    documents.extend(pdf_docs)
                    
                elif file_ext == '.docx' and DOCX_AVAILABLE:
                    docx_docs = self._load_docx(file_path)
                    documents.extend(docx_docs)
                    
                elif file_ext in {'.md', '.markdown'}:
                    md_docs = self._load_markdown(file_path)
                    documents.extend(md_docs)
                    
            except Exception as e:
                logger.error(f"Error loading {file_path}: {e}")
                continue
        
        return documents
    
    def _load_docx(self, file_path: Path) -> List[Document]:
        """
        Load DOCX file and convert to Document objects with section/page metadata.
        """
        documents = []
        
        try:
            doc = DocxDocument(str(file_path))
            file_name = file_path.name
            
            # Extract text by paragraphs (for section tracking)
            paragraphs = []
            current_section = 1
            section_text = []
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                
                # Detect section headers (bold, larger font, or specific patterns)
                is_header = (
                    para.style.name.startswith('Heading') or
                    para.style.name.startswith('Title') or
                    para.runs and any(run.bold for run in para.runs) or
                    len(text) < 100 and text.isupper()
                )
                
                if is_header and section_text:
                    # Save previous section
                    section_text_combined = '\n'.join(section_text)
                    if section_text_combined.strip():
                        doc_obj = Document(
                            text=section_text_combined,
                            metadata={
                                'file_name': file_name,
                                'page_label': str(current_section),  # Use section number as page_label
                                'content_type': 'text',
                                'file_type': 'docx',
                                'section_number': current_section
                            }
                        )
                        documents.append(doc_obj)
                        current_section += 1
                        section_text = []
                
                section_text.append(text)
            
            # Add final section
            if section_text:
                section_text_combined = '\n'.join(section_text)
                if section_text_combined.strip():
                    doc_obj = Document(
                        text=section_text_combined,
                        metadata={
                            'file_name': file_name,
                            'page_label': str(current_section),
                            'content_type': 'text',
                            'file_type': 'docx',
                            'section_number': current_section
                        }
                    )
                    documents.append(doc_obj)
            
            # If no sections detected, create single document
            if not documents:
                full_text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
                if full_text.strip():
                    doc_obj = Document(
                        text=full_text,
                        metadata={
                            'file_name': file_name,
                            'page_label': '1',
                            'content_type': 'text',
                            'file_type': 'docx',
                            'section_number': 1
                        }
                    )
                    documents.append(doc_obj)
            
            logger.info(f"Loaded DOCX {file_name}: {len(documents)} sections")
            
        except Exception as e:
            logger.error(f"Error loading DOCX {file_path}: {e}")
        
        return documents
    
    def _load_markdown(self, file_path: Path) -> List[Document]:
        """
        Load Markdown file and convert to Document objects with section metadata.
        """
        documents = []
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            file_name = file_path.name
            
            # Split by markdown headers (# Header, ## Header, etc.)
            lines = content.split('\n')
            sections = []
            current_section = []
            current_section_num = 1
            
            for line in lines:
                # Check if line is a markdown header
                is_header = False
                header_level = 0
                
                if line.strip().startswith('#'):
                    # Count # symbols
                    header_level = len(line) - len(line.lstrip('#'))
                    if header_level <= 3:  # Only treat h1-h3 as section breaks
                        is_header = True
                
                if is_header and current_section:
                    # Save previous section
                    section_text = '\n'.join(current_section).strip()
                    if section_text:
                        sections.append({
                            'text': section_text,
                            'section_num': current_section_num,
                            'header': line.strip()
                        })
                        current_section_num += 1
                    current_section = []
                
                current_section.append(line)
            
            # Add final section
            if current_section:
                section_text = '\n'.join(current_section).strip()
                if section_text:
                    sections.append({
                        'text': section_text,
                        'section_num': current_section_num,
                        'header': ''
                    })
            
            # Create Document objects
            for section in sections:
                doc_obj = Document(
                    text=section['text'],
                    metadata={
                        'file_name': file_name,
                        'page_label': str(section['section_num']),  # Use section number as page_label
                        'content_type': 'text',
                        'file_type': 'markdown',
                        'section_number': section['section_num'],
                        'section_header': section['header']
                    }
                )
                documents.append(doc_obj)
            
            # If no sections detected, create single document
            if not documents and content.strip():
                doc_obj = Document(
                    text=content,
                    metadata={
                        'file_name': file_name,
                        'page_label': '1',
                        'content_type': 'text',
                        'file_type': 'markdown',
                        'section_number': 1
                    }
                )
                documents.append(doc_obj)
            
            logger.info(f"Loaded Markdown {file_name}: {len(documents)} sections")
            
        except Exception as e:
            logger.error(f"Error loading Markdown {file_path}: {e}")
        
        return documents


class TechnicalRAGPipeline:
    """High-performance RAG pipeline optimized for technical documentation with non-text content support."""
    
    def __init__(self, cache_dir="/root/.cache/huggingface/hub", config_path="config.yaml"):
        self.cache_dir = cache_dir
        self.embed_model = None
        self.reranker = None
        self.index = None
        self.non_text_extractor = NonTextExtractor()
        self.text_preprocessor = TextPreprocessor()
        self.config = self._load_config(config_path)
        
        # Initialize Claude semantic rewriter (optional)
        claude_config = self.config.get("claude_rewriting", {})
        self.claude_rewriter = ClaudeSemanticRewriter(
            api_key=claude_config.get("api_key"),
            model=claude_config.get("model", "claude-3-5-sonnet-20241022"),
            enabled=claude_config.get("enabled", False),
            max_retries=claude_config.get("max_retries", 2),
            timeout=claude_config.get("timeout", 30)
        )
        
    def _load_config(self, config_path: str) -> Dict[str, Any]:
        """Load configuration from YAML file or use defaults."""
        default_config = {
            "qdrant": {
                "url": "http://localhost:6333",
                "collection_name": "technical_docs"
            },
            "models": {
                "embedding": "BAAI/bge-large-en-v1.5",
                "reranker": "BAAI/bge-reranker-large"
            },
            "chunking": {
                "chunk_size": 512,
                "chunk_overlap": 128
            }
        }
        
        if os.path.exists(config_path):
            try:
                with open(config_path, 'r') as f:
                    config = yaml.safe_load(f)
                return {**default_config, **config}
            except Exception as e:
                logger.warning(f"Failed to load config: {e}, using defaults")
        
        return default_config
        
    def initialize_models(self):
        """Initialize embedding and re-ranking models."""
        logger.info("🚀 Initializing embedding model...")
        
        # Disable hf_transfer if not installed (RunPod issue)
        import os
        import shutil
        if os.environ.get('HF_HUB_ENABLE_HF_TRANSFER') == '1':
            logger.info("Disabling HF_HUB_ENABLE_HF_TRANSFER (package not installed)")
            os.environ['HF_HUB_ENABLE_HF_TRANSFER'] = '0'
        
        # Detect GPU with fallback to CPU if CUDA is incompatible
        import torch
        device = "cpu"  # Default to CPU
        try:
            if torch.cuda.is_available():
                # Test if CUDA actually works with a real operation
                try:
                    test_tensor = torch.zeros(1).cuda()
                    # Try a simple operation that requires kernel execution
                    result = test_tensor + 1
                    result.item()  # Force execution
                    del test_tensor, result
                    torch.cuda.empty_cache()
                    device = "cuda"
                    logger.info(f"🖥️ Using device: {device}")
                    logger.info(f"GPU: {torch.cuda.get_device_name(0)}")
                    logger.info(f"GPU Memory: {torch.cuda.get_device_properties(0).total_memory / 1e9:.2f} GB")
                except RuntimeError as cuda_error:
                    if "CUDA" in str(cuda_error) or "kernel" in str(cuda_error).lower():
                        logger.warning(f"⚠️ CUDA not compatible (kernel error), falling back to CPU: {cuda_error}")
                        device = "cpu"
                        logger.info(f"🖥️ Using device: {device}")
                    else:
                        raise
            else:
                logger.info(f"🖥️ Using device: {device} (CUDA not available)")
        except Exception as e:
            logger.warning(f"⚠️ CUDA test failed, falling back to CPU: {e}")
            device = "cpu"
            logger.info(f"🖥️ Using device: {device}")
        
        cache_path = os.path.expanduser(self.cache_dir)
        
        # Try multiple approaches
        model_options = [
            ("BAAI/bge-large-en-v1.5", "BGE Large"),
            ("BAAI/bge-base-en-v1.5", "BGE Base"),
            ("all-MiniLM-L6-v2", "MiniLM"),
            ("all-mpnet-base-v2", "MPNet")
        ]
        
        for model_name, display_name in model_options:
            try:
                logger.info(f"Trying model: {display_name} ({model_name})")
                
                # Method 1: Direct load without sentence-transformers prefix
                try:
                    # Force CPU if CUDA is incompatible
                    if device == "cuda":
                        try:
                            # Test CUDA compatibility
                            import torch
                            test = torch.zeros(1).cuda()
                            del test
                        except Exception:
                            logger.warning(f"CUDA incompatible, forcing CPU for {display_name}")
                            device = "cpu"
                    
                    self.embed_model = HuggingFaceEmbedding(
                        model_name=model_name,
                        cache_folder=self.cache_dir,
                        trust_remote_code=True,
                        device=device
                    )
                    logger.info(f"✅ Successfully loaded: {display_name} on {device}")
                    break
                except Exception as e1:
                    logger.debug(f"Method 1 failed: {e1}")
                    
                    # Method 2: Try with full sentence-transformers path
                    if not model_name.startswith("sentence-transformers/"):
                        try:
                            full_name = f"sentence-transformers/{model_name}"
                            self.embed_model = HuggingFaceEmbedding(
                                model_name=full_name,
                                cache_folder=self.cache_dir,
                                trust_remote_code=True,
                                device=device
                            )
                            logger.info(f"✅ Successfully loaded: {display_name} on {device}")
                            break
                        except Exception as e2:
                            logger.debug(f"Method 2 failed: {e2}")
                            raise e1
                    else:
                        raise e1
                        
            except Exception as e:
                logger.warning(f"Failed to load {display_name}: {str(e)[:100]}")
                continue
        
        if not self.embed_model:
            logger.error("All model loading attempts failed. Trying emergency fallback...")
            # Emergency fallback - use any available model
            try:
                self.embed_model = HuggingFaceEmbedding(model_name="all-MiniLM-L6-v2")
                logger.info("✅ Loaded with emergency fallback")
            except:
                raise RuntimeError("Could not load any embedding model. Check internet connection and HuggingFace access.")
        
        # Try to initialize re-ranker (optional)
        try:
            logger.info("🎯 Initializing re-ranker...")
            reranker_model = self.config.get("models", {}).get("reranker", "BAAI/bge-reranker-large")
            # Use CPU if CUDA was incompatible
            reranker_device = device if device == "cpu" else device
            try:
                self.reranker = CrossEncoder(
                    reranker_model,
                    cache_folder=self.cache_dir,
                    device=reranker_device
                )
                logger.info(f"✅ Re-ranker loaded successfully on {reranker_device}")
            except RuntimeError as cuda_error:
                if "CUDA" in str(cuda_error) or "cuda" in str(cuda_error).lower():
                    logger.warning(f"⚠️ CUDA incompatible for reranker, using CPU: {cuda_error}")
                    self.reranker = CrossEncoder(
                        reranker_model,
                        cache_folder=self.cache_dir,
                        device="cpu"
                    )
                    logger.info(f"✅ Re-ranker loaded successfully on CPU")
                else:
                    raise
        except Exception as e:
            logger.warning(f"Re-ranker not available: {e}")
            self.reranker = None
        
        # Set global embedding model
        Settings.embed_model = self.embed_model
        logger.info("✅ Models initialized successfully")
    
    def process_non_text_content(self, data_dir: str) -> Tuple[List[Dict], List[Dict], List[Dict]]:
        """Process non-text content (tables, images, captions) from documents."""
        logger.info("📊 Processing non-text content...")
        
        all_tables = []
        all_images = []
        all_captions = []
        
        # Find all PDF files
        pdf_files = list(Path(data_dir).glob("*.pdf"))
        logger.info(f"Found {len(pdf_files)} PDF files to process")
        
        for pdf_path in pdf_files:
            logger.info(f"Processing {pdf_path.name}...")
            
            try:
                # Extract tables
                tables = self.non_text_extractor.extract_tables_from_pdf(str(pdf_path))
                all_tables.extend(tables)
                logger.info(f"Extracted {len(tables)} tables from {pdf_path.name}")
                
                # Extract images
                images = self.non_text_extractor.extract_images_from_pdf(str(pdf_path))
                all_images.extend(images)
                logger.info(f"Extracted {len(images)} images from {pdf_path.name}")
                
                # Extract captions
                captions = self.non_text_extractor.extract_figure_captions(str(pdf_path))
                all_captions.extend(captions)
                logger.info(f"Extracted {len(captions)} captions from {pdf_path.name}")
                
            except Exception as e:
                logger.error(f"Failed to process {pdf_path.name}: {e}")
                continue
        
        logger.info(f"✅ Non-text processing complete: {len(all_tables)} tables, {len(all_images)} images, {len(all_captions)} captions")
        return all_tables, all_images, all_captions
    
    def create_non_text_nodes(self, tables: List[Dict], images: List[Dict], captions: List[Dict]) -> List[TextNode]:
        """Create TextNode objects for non-text content to be embedded."""
        nodes = []
        
        # Process tables
        for table in tables:
            # Create text representation of table
            table_text = f"Table from {Path(table['source_path']).name}, page {table['page_number']}:\n{table['table_markdown']}"
            
            node = TextNode(
                text=table_text,
                metadata={
                    "content_type": "table",
                    "source_path": table["source_path"],
                    "page_number": table["page_number"],
                    "table_index": table["table_index"],
                    "row_count": table["row_count"],
                    "column_count": table["column_count"],
                    "table_json": table["table_json"]
                }
            )
            nodes.append(node)
        
        # Process figure captions
        for caption in captions:
            # Clean caption text to remove boilerplate
            caption_text = self.text_preprocessor.clean_text(caption["caption_text"])
            if not self.text_preprocessor.should_skip_node(caption_text):
                node = TextNode(
                    text=caption_text,
                    metadata={
                        "content_type": "figure_caption",
                        "source_path": caption["source_path"],
                        "page_number": caption["page_number"],
                        "line_number": caption["line_number"]
                    }
                )
                nodes.append(node)
        
        # Process images (create text nodes for captions and metadata only - no base64)
        for image in images:
            image_text = f"Image from {Path(image['source_path']).name}, page {image['page_number']}: {image['caption']}"
            
            node = TextNode(
                text=image_text,
                metadata={
                    "content_type": "image",
                    "source_path": image["source_path"],
                    "page_number": image["page_number"],
                    "image_index": image["image_index"],
                    "width": image["width"],
                    "height": image["height"],
                    "saved_path": image.get("saved_path"),
                    "bbox": str(image.get("bbox")) if image.get("bbox") else None
                }
            )
            nodes.append(node)
        
        return nodes
    
    def setup_qdrant_storage(self) -> StorageContext:
        """Setup Qdrant vector store for hybrid search."""
        try:
            # Initialize Qdrant client
            qdrant_url = self.config["qdrant"]["url"]
            collection_name = self.config["qdrant"]["collection_name"]
            
            client = qdrant_client.QdrantClient(url=qdrant_url)
            
            # Create vector store
            vector_store = QdrantVectorStore(
                client=client,
                collection_name=collection_name
            )
            
            # Create storage context
            storage_context = StorageContext.from_defaults(
                vector_store=vector_store,
                docstore=SimpleDocumentStore(),
                index_store=SimpleIndexStore()
            )
            
            logger.info(f"✅ Qdrant storage configured: {qdrant_url}/{collection_name}")
            return storage_context
            
        except Exception as e:
            logger.warning(f"Qdrant not available, using local storage: {e}")
            return None
    
    def build_index(self, data_dir="data", storage_dir="latest_model", use_qdrant=False):
        """Build or load vector index with optimized chunking and non-text content."""
        
        # Initialize models
        self.initialize_models()
        
        # Setup storage context
        storage_context = None
        if use_qdrant:
            storage_context = self.setup_qdrant_storage()
        
        # For local storage: always rebuild (clear old index first)
        if not use_qdrant and os.path.exists(storage_dir):
            logger.info(f"🗑️  Clearing old index from {storage_dir} for fresh rebuild...")
            shutil.rmtree(storage_dir)
            logger.info("✅ Old index cleared - ready for fresh build")
        
        # Create storage directory if it doesn't exist
        if not use_qdrant:
            os.makedirs(storage_dir, exist_ok=True)
        
        print("\n" + "="*70)
        print("📥 BUILDING NEW RAG INDEX")
        print("="*70)
        
        # Step 1: Load Documents (PDF, DOCX, Markdown)
        print("\n[Step 1/7] 📄 Loading documents (PDF, DOCX, Markdown)...")
        loader = DocumentLoader(data_dir)
        documents = loader.load_documents()
        
        # Count by file type
        file_types = {}
        for doc in documents:
            file_type = doc.metadata.get('file_type', 'pdf')
            file_types[file_type] = file_types.get(file_type, 0) + 1
        
        type_summary = ", ".join([f"{count} {ftype.upper()}" for ftype, count in file_types.items()])
        print(f"   ✅ Loaded {len(documents)} document sections ({type_summary})")
        logger.info(f"Loaded {len(documents)} documents: {type_summary}")
        
        # Step 2: Enhanced AI-Powered Preprocessing
        print("\n[Step 2/7] 🧹 Enhanced preprocessing (TOC removal, artifact fixing, normalization)...")
        preprocessed_docs = []
        skipped_pages = 0
        skip_reasons = {}
        
        for doc in documents:
            original_text = doc.text or ""
            
            # Enhanced cleaning with metadata for context-aware processing
            cleaned_text = self.text_preprocessor.clean_text(original_text, metadata=doc.metadata)
            
            # Check if page should be skipped (with reason tracking)
            if self.text_preprocessor.is_low_content_page(cleaned_text):
                skip_reasons['low_content'] = skip_reasons.get('low_content', 0) + 1
                skipped_pages += 1
                logger.debug(f"Skipping low-content page: {doc.metadata.get('file_name', 'unknown')}")
                continue
            
            # Check for first page without content
            if self.text_preprocessor.is_first_page_without_content(cleaned_text, metadata=doc.metadata):
                skip_reasons['first_page_no_content'] = skip_reasons.get('first_page_no_content', 0) + 1
                skipped_pages += 1
                logger.debug(f"Skipping first page without content: {doc.metadata.get('file_name', 'unknown')}")
                continue
            
            # Create new document with cleaned text
            if cleaned_text:  # Only add if there's content left after cleaning
                new_doc = Document(
                    text=cleaned_text,
                    metadata=doc.metadata
                )
                preprocessed_docs.append(new_doc)
        
        skip_summary = ", ".join([f"{reason}: {count}" for reason, count in skip_reasons.items()]) if skip_reasons else "none"
        print(f"   ✅ Preprocessed {len(preprocessed_docs)} documents ({skipped_pages} pages skipped: {skip_summary})")
        logger.info(f"Preprocessed {len(preprocessed_docs)} documents, skipped {skipped_pages} pages ({skip_summary})")
        
        # Step 3: Extract Non-Text Content
        print("\n[Step 3/7] 🖼️  Extracting tables, images, and captions...")
        print("   This may take a few minutes...")
        tables, images, captions = self.process_non_text_content(data_dir)
        print(f"   ✅ Extracted {len(tables)} tables, {len(images)} images, {len(captions)} captions")
        
        # Step 4: Create Non-Text Nodes
        print("\n[Step 4/7] 📊 Creating searchable nodes from extracted content...")
        non_text_nodes = self.create_non_text_nodes(tables, images, captions)
        print(f"   ✅ Created {len(non_text_nodes)} non-text nodes")
        logger.info(f"Created {len(non_text_nodes)} non-text nodes")
        
        # Step 5: Smart Chunking with Text Nodes
        print("\n[Step 5/7] 🧠 Smart chunking and filtering...")
        chunk_size = self.config.get("chunking", {}).get("chunk_size", 350)
        chunk_overlap = self.config.get("chunking", {}).get("chunk_overlap", 88)
        
        smart_splitter = SmartChunkSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            preprocessor=self.text_preprocessor
        )
        
        print(f"   - Chunk size: {chunk_size} characters")
        print(f"   - Chunk overlap: {chunk_overlap} characters")
        print(f"   - Preserving structured content (tables, code blocks, numbered steps)...")
        
        # Get nodes from documents using smart chunking
        text_nodes = smart_splitter.get_nodes_from_documents(preprocessed_docs, show_progress=True)
        
        # Filter out short/low-quality nodes (only for text nodes, not tables/images)
        filtered_nodes = []
        skipped_nodes = 0
        skip_reasons_chunks = {}
        
        for node in text_nodes:
            # Skip filtering for non-text content types (tables, images, captions are already handled separately)
            content_type = node.metadata.get("content_type", "text")
            if content_type != "text":
                filtered_nodes.append(node)  # Don't filter non-text content
            else:
                # Enhanced skip check with reason tracking
                should_skip, skip_reason = self.text_preprocessor.should_skip_node(
                    node.text, 
                    metadata=node.metadata
                )
                if should_skip:
                    skipped_nodes += 1
                    skip_reasons_chunks[skip_reason] = skip_reasons_chunks.get(skip_reason, 0) + 1
                    logger.debug(f"Skipping node from {node.metadata.get('file_name', 'unknown')}: {skip_reason}")
                else:
                    filtered_nodes.append(node)
        
        skip_summary_chunks = ", ".join([f"{reason}: {count}" for reason, count in skip_reasons_chunks.items()]) if skip_reasons_chunks else "none"
        print(f"   ✅ Created {len(filtered_nodes)} text nodes ({skipped_nodes} filtered: {skip_summary_chunks})")
        logger.info(f"Created {len(filtered_nodes)} text nodes, filtered {skipped_nodes} nodes ({skip_summary_chunks})")
        
        # Step 5.5: Optional Claude Semantic Rewriting
        if self.claude_rewriter.enabled:
            print("\n[Step 5.5/7] 🤖 Claude semantic rewriting (improving clarity while preserving meaning)...")
            print("   - This step uses Claude API to enhance text clarity")
            print("   - Structured content (tables, code, lists) will be preserved as-is")
            print("   - Estimated time: 1-3 minutes per 100 chunks")
            
            rewritten_nodes, rewrite_stats = self.claude_rewriter.rewrite_nodes(filtered_nodes, show_progress=True)
            
            print(f"   ✅ Rewriting complete:")
            print(f"      - Rewritten: {rewrite_stats['rewritten']} chunks")
            print(f"      - Preserved (structured): {rewrite_stats['structured']} chunks")
            print(f"      - Skipped: {rewrite_stats['skipped']} chunks")
            print(f"      - Failed (using original): {rewrite_stats['failed']} chunks")
            
            # Use rewritten nodes for embedding
            filtered_nodes = rewritten_nodes
        else:
            logger.debug("Claude rewriting disabled, skipping rewrite step")
        
        # Step 6: Create Vector Embeddings (LONGEST STEP)
        print("\n[Step 6/7] 🧠 Generating embeddings and building vector index...")
        print(f"   - Processing {len(filtered_nodes)} text nodes + {len(non_text_nodes)} non-text nodes...")
        print(f"   - This is the LONGEST step (embedding generation)")
        print(f"   - Expected time: 5-15 minutes on GPU, 30-60 minutes on CPU")
        print(f"   - Watch for progress below...")
        print("")
        
        # Record start time
        import time
        start_time = time.time()
        
        # Combine all nodes
        all_nodes = filtered_nodes + non_text_nodes
        
        # Create index from nodes (LlamaIndex API)
        if storage_context:
            # Create index with storage context
            self.index = VectorStoreIndex(
                nodes=[],
                storage_context=storage_context,
                show_progress=True
            )
        else:
            # Create index without storage context (will use default)
            self.index = VectorStoreIndex(
                nodes=[],
                show_progress=True
            )
        
        # Insert all nodes into the index (batch insert for better performance)
        print(f"   Inserting {len(all_nodes)} nodes into index...")
        batch_size = 100  # Insert in batches
        for i in tqdm(range(0, len(all_nodes), batch_size), desc="   Inserting nodes", unit="batch"):
            batch = all_nodes[i:i + batch_size]
            try:
                self.index.insert_nodes(batch)
            except RuntimeError as e:
                if "CUDA" in str(e) or "cuda" in str(e).lower():
                    logger.error(f"CUDA error during embedding - device should have been set to CPU")
                    logger.error(f"Error: {e}")
                    raise RuntimeError("CUDA incompatible. The embedding model should use CPU. Check device detection.") from e
                else:
                    raise
        
        elapsed = time.time() - start_time
        print(f"\n   ✅ Vector index created in {elapsed:.1f} seconds ({elapsed/60:.1f} minutes)")
        print(f"   ⚡ Processing speed: {len(all_nodes) / elapsed:.2f} nodes/sec")
        
        # Persist the index (only for local storage)
        print(f"\n💾 Saving index to disk...")
        if not use_qdrant:
            self.index.storage_context.persist(persist_dir=storage_dir)
            print(f"   ✅ Index saved to: {storage_dir}")
            logger.info("✅ Index created and saved locally")
        else:
            print(f"   ✅ Index saved to: Qdrant")
            logger.info("✅ Index created and saved to Qdrant")
        
        # Final summary
        total_time = time.time() - start_time
        print("\n" + "="*70)
        print("✅ INGESTION COMPLETE!")
        print("="*70)
        print(f"📊 Documents loaded: {len(documents)}")
        print(f"📊 Documents after preprocessing: {len(preprocessed_docs)} ({skipped_pages} low-content pages skipped)")
        print(f"📊 Text nodes created: {len(filtered_nodes)} ({skipped_nodes} filtered)")
        print(f"📊 Non-text nodes: {len(non_text_nodes)}")
        print(f"📊 Total nodes indexed: {len(all_nodes)}")
        if self.claude_rewriter.enabled:
            print(f"🤖 Claude rewriting: Enabled (check logs for rewrite statistics)")
        print(f"⏱️  Total time: {total_time:.1f} seconds ({total_time/60:.1f} minutes)")
        if not use_qdrant:
            print(f"📁 Storage location: {storage_dir}")
            print(f"\n💡 Two-Pod Workflow:")
            print(f"   1. git add {storage_dir}/")
            print(f"   2. git commit -m 'Update RAG index'")
            print(f"   3. git push")
            print(f"   4. Switch to cheap pod → git pull → streamlit run app.py")
        print(f"📂 Extracted content: extracted_content/")
        print(f"🔍 Ready to query!")
        print("="*70 + "\n")
        
        return self.index
    
    def hybrid_search(self, query: str, top_k: int = 10, content_types: List[str] = None) -> List[NodeWithScore]:
        """Perform hybrid search across text, tables, and images."""
        if not self.index:
            raise RuntimeError("Index not built. Call build_index() first.")
        
        # Default to search all content types
        if content_types is None:
            content_types = ["text", "table", "image", "figure_caption"]
        
        # Perform vector search
        retriever = self.index.as_retriever(similarity_top_k=top_k * 2)  # Get more for re-ranking
        nodes = retriever.retrieve(query)
        
        # Filter by content type if specified
        if content_types:
            filtered_nodes = []
            for node in nodes:
                content_type = node.metadata.get("content_type", "text")
                if content_type in content_types or content_type == "text":
                    filtered_nodes.append(node)
            nodes = filtered_nodes[:top_k]
        
        # Apply re-ranking if available
        if self.reranker and len(nodes) > 1:
            logger.info("🎯 Applying re-ranking...")
            try:
                # Prepare query-document pairs for re-ranking
                pairs = [(query, node.text) for node in nodes]
                scores = self.reranker.predict(pairs)
                
                # Sort by re-ranking scores
                scored_nodes = list(zip(nodes, scores))
                scored_nodes.sort(key=lambda x: x[1], reverse=True)
                nodes = [node for node, score in scored_nodes[:top_k]]
                
            except Exception as e:
                logger.warning(f"Re-ranking failed: {e}")
        
        return nodes[:top_k]
    
    def backup_storage(self, storage_dir: str, backup_name: str = None):
        """Create a backup of the storage directory."""
        if not os.path.exists(storage_dir):
            logger.warning(f"Storage directory {storage_dir} does not exist")
            return None
            
        if backup_name is None:
            backup_name = f"rag_backup_{int(time.time())}"
        
        backup_path = f"/workspace/{backup_name}.tar.gz"
        
        try:
            with tarfile.open(backup_path, "w:gz") as tar:
                tar.add(storage_dir, arcname=os.path.basename(storage_dir))
            
            logger.info(f"✅ Backup created: {backup_path}")
            return backup_path
        except Exception as e:
            logger.error(f"Failed to create backup: {e}")
            return None
    
    def restore_storage(self, backup_path: str, storage_dir: str):
        """Restore storage from backup."""
        try:
            with tarfile.open(backup_path, "r:gz") as tar:
                tar.extractall(os.path.dirname(storage_dir))
            logger.info(f"✅ Storage restored from {backup_path}")
            return True
        except Exception as e:
            logger.error(f"Failed to restore backup: {e}")
            return False


def main():
    """Main function to build the RAG index with non-text content support."""
    
    # Initialize pipeline
    pipeline = TechnicalRAGPipeline()
    
    # Build or load index (set use_qdrant=True for Qdrant storage)
    use_qdrant = os.getenv("USE_QDRANT", "false").lower() == "true"
    index = pipeline.build_index(use_qdrant=use_qdrant)
    
    print("\n" + "="*60)
    print("✅ INGESTION COMPLETED SUCCESSFULLY")
    print("="*60)
    if use_qdrant:
        print("🗄️ Index saved to: Qdrant")
    else:
        print("📁 Index saved to: storage/")
    print("🔍 Use query.py to search the documents")
    print("📊 Non-text content extracted to: extracted_content/")
    print("="*60)


if __name__ == "__main__":
    main()